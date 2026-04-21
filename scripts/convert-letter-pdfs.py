#!/usr/bin/env python3
"""
Konvertiert die 10 Standardschreiben-PDFs nach DOCX via Claude Vision.

Ansatz:
  Für jede PDF-Seite rendert das Skript ein hochauflösendes Bild und schickt
  es an Claude (Sonnet). Claude gibt strukturiertes JSON zurück: pro Absatz
  eine Liste von Spans mit text + bold/italic-Flags. python-docx schreibt
  daraus ein Blocksatz-DOCX mit korrekten Bold/Italic-Runs und Leerabsätzen
  zwischen logischen Blöcken.

  Warum Claude Vision:
  - Die PDFs verwenden eine einzige Font ("MaraDocsFont", flags=12), Bold
    wird per Stroke simuliert → pymupdf erkennt kein Bold-Flag
  - Claude sieht das Layout visuell, merged mehrzeilige Sätze korrekt,
    erkennt Kursiv/Fett zuverlässig und ignoriert OCR-Artefakte

Platzhalter-Normalisierung (nach Claude-Output):
  FELD_Schuldners-Schuldnerin → FELD_Schuldners_Schuldnerin, etc.
  Alt-Syntax im Steuerberater-Brief: [Datum] → FELD_Akte_LastGAW, u.a.

Usage:
  python3 scripts/convert-letter-pdfs.py
  python3 scripts/convert-letter-pdfs.py --only bausparkassen-anfrage.docx
"""

from __future__ import annotations
import argparse
import base64
import json
import os
import re
import sys
from pathlib import Path
from typing import Any

import fitz  # pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from dotenv import load_dotenv

REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "standardschreiben"
DST = SRC / "templates"

load_dotenv(REPO / ".env")

import anthropic

PDF_MAP = {
    "Bankenanfrage.pdf": "Bankenanfrage.docx",
    "Anfrage_zu_bestehendem_Vertragsverhältnis_Bausparkasse_.pdf": "Anfrage_zu_bestehendem_Vertragsverhältnis_Bausparkasse_.docx",
    "Muster_Kontaktaufnahme_Steuerberater.pdf": "Muster_Kontaktaufnahme_Steuerberater.docx",
    "Einsichtnahmegesuch_Strafakte_Anfrage_zur_Akteneinsicht_.pdf": "Einsichtnahmegesuch_Strafakte_Anfrage_zur_Akteneinsicht_.docx",
    "Halteranfrage_Zulassungsstelle.pdf": "Halteranfrage_Zulassungsstelle.docx",
    "Gewerbeanfrage.pdf": "Gewerbeanfrage.docx",
    "Anfrage_ans_Finanzamt.pdf": "Anfrage_ans_Finanzamt.docx",
    "Halteranfrage_Kraftfahrt_Bundesamt.pdf": "Halteranfrage_Kraftfahrt_Bundesamt.docx",
    "Muster_Versicherungsanfrage.pdf": "Muster_Versicherungsanfrage.docx",
    "Gerichtsvollzieheranfrage.pdf": "Gerichtsvollzieheranfrage.docx",
}

# Manuell nachbearbeitete Vorlagen — nicht überschreiben
SKIP_FILES = {
    "Bankenanfrage.docx",
    "Anfrage_ans_Finanzamt.docx",
    "Anfrage_zu_bestehendem_Vertragsverhältnis_Bausparkasse_.docx",
    "Einsichtnahmegesuch_Strafakte_Anfrage_zur_Akteneinsicht_.docx",
    "Halteranfrage_Zulassungsstelle.docx",
    "Halteranfrage_Kraftfahrt_Bundesamt.docx",
}

KNOWN_FIELDS = {
    "FELD_Akte_Aktenzeichen", "FELD_Akte_Gericht",
    "FELD_Akte_LastGAVV",  # Beschlussdatum (TBS-Konvention)
    "FELD_Akte_EroeffDat", "FELD_Akte_Bezeichnung", "FELD_Akte_VerfahrenArt",
    "FELD_Gericht_Ort",
    "FELD_Schuldner_Name", "FELD_Schuldner_Vorname", "FELD_Schuldner_Vollname",
    "FELD_Schuldner_Adr", "FELD_Schuldner_Adresse", "FELD_Schuldner_Geburtsdatum",
    "FELD_Schuldner_Firma", "FELD_Schuldner_Betriebsstaette", "FELD_Schuldner_HRB",
    "FELD_Schuldner_Artikel", "FELD_Schuldner_der_die", "FELD_Schuldner_Der_Die_Groß",
    "FELD_Schuldner_den_die", "FELD_Schuldner_dem_der",
    "FELD_Schuldner_Schuldnerin", "FELD_Schuldners_Schuldnerin",
    "FELD_Schuldner_Halters_Halterin",
    "FELD_Strafverfahren_Person", "FELD_Strafverfahren_Tatvorwurf", "FELD_Strafverfahren_Gegenstand",
    "FELD_Verwalter_Name", "FELD_Verwalter_Art", "FELD_Verwalter_Diktatzeichen",
    "FELD_Verwalter_Unterzeichner", "FELD_Verwalter_zum_zur", "FELD_Verwalter_der_die",
    "FELD_Verwalter_Der_Die_Groß",
    "FELD_Bet_AnredeHoeflichOV", "FELD_Bet_GrussBriefende",
    "FELD_ANSCHREIBEN_DAT_2",
}

POST_FIXES = [
    (re.compile(r"FELD__+"), "FELD_"),
    (re.compile(r"(FELD_[A-Za-zÄÖÜäöüß0-9]+)-([A-Za-zÄÖÜäöüß0-9][A-Za-zÄÖÜäöüß0-9_]*)"), r"\1_\2"),
    (re.compile(r"(FELD_[A-Za-zÄÖÜäöüß0-9]+) _([A-Za-zÄÖÜäöüß0-9][A-Za-zÄÖÜäöüß0-9_]*)"), r"\1_\2"),
    # Beschlussdatum-Konvention: GAW und GAV werden als GAVV normalisiert
    (re.compile(r"FELD_Akte_LastGAW\b"), "FELD_Akte_LastGAVV"),
    (re.compile(r"FELD_Akte_LastGAV\b"), "FELD_Akte_LastGAVV"),
    (re.compile(r"\bAnlage Il\b"), "Anlage II"),
    (re.compile(r"\bAnlage IIl\b"), "Anlage III"),
]

PER_FILE_FIXES = {
    "Muster_Kontaktaufnahme_Steuerberater.docx": [
        (re.compile(r"\[Name der Schuldnerin/des Schuldners\]"), "FELD_Schuldner_Vollname"),
        (re.compile(r"\[Datum\]"), "FELD_Akte_LastGAW"),
        # "[Funktionsbezeichnung, z. B. als Sachverständige/r / vorläufige/r Insolvenzverwalter/in]"
        (re.compile(r"\[Funktionsbezeichnung[^\]]*\]"), "FELD_Verwalter_Art"),
        # "[Name]" im Signatur-Block → Verwalter, nicht Schuldner
        (re.compile(r"\[Name\]"), "FELD_Verwalter_Name"),
    ],
}

# Nach Generierung: Spans, bei denen der Text mit einem dieser Strings beginnt,
# werden auf bold=False zurückgesetzt (Claude markiert sie fälschlich als bold).
FORCE_PLAIN_PREFIXES = {
    "Einsichtnahmegesuch_Strafakte_Anfrage_zur_Akteneinsicht_.docx": [
        "mir Einsicht in die betreffende Strafakte",
    ],
}

SYSTEM_PROMPT = """Du extrahierst Text aus Vorlagen-Briefen (Standardschreiben) und gibst ihn als
strukturiertes JSON zurück.

Die Briefe enthalten deutsche Fachsprache und Platzhalter der Form `FELD_Irgendwas_Name`,
die wörtlich übernommen werden müssen.

Gib ein JSON-Objekt zurück mit:
{
  "paragraphs": [
    {
      "spans": [
        {"text": "string", "bold": bool, "italic": bool, "underline": bool}
      ]
    }
  ]
}

KRITISCHE REGELN:

1. **Absätze**: Ein Absatz pro logischer Texteinheit. Mehrzeilige Sätze im Bild →
   EIN Absatz, Zeilenumbrüche innerhalb eines Satzes werden ignoriert. Zwischen
   verschiedenen Absätzen wird automatisch ein Leerabsatz eingefügt — füge KEINE
   leeren Einträge ein.

2. **Underline (sehr wichtig)**: Die Vorlagen verwenden UNTERSTRICHENEN Text, um
   wichtige Passagen hervorzuheben. Unterstrichener Text wird als
   `"underline": true` markiert und bleibt Unterstreichung (nicht zu Bold konvertieren).
   Typische Beispiele:
     - "beigefügten Beschluss" oder ähnliche Verweise auf Anlagen
     - "Anlage II", "Anlage III" (Verweise auf Anlagen)
     - "Anlagen:" / "Anlage:" (Listen-Header)
     - Ganze Sätze mit Handlungsaufforderungen
     - Fristen und Datumsangaben
   Prüfe sehr genau: Unterstreichung kann dünn sein. Im Zweifel Text als underline
   markieren, aber nur wenn du eine Linie unter dem Text siehst.

3. **Bold**: `"bold": true` NUR bei tatsächlich fetten Buchstaben (dickere Striche
   als umgebender Text). Diese Vorlagen haben selten echtes Bold — meist ist
   das, was fett wirkt, eigentlich Unterstreichung.
   WICHTIG: Betreff-Zeilen ("Insolvenzantragsverfahren ...") und eingerückte
   Hervorhebungen sind NICHT bold, nur optisch ausgerückt. Markiere diese als
   `"bold": false, "underline": false`, es sei denn du siehst eindeutig dickere
   Striche oder eine Unterstreichungslinie.

4. **Italic**: `"italic": true` nur bei geneigten Buchstaben (Kursiv).

5. **Mehrere Spans pro Absatz**: Wenn in einem Absatz nur ein Teil hervorgehoben
   ist, teile den Absatz in mehrere Spans auf:
     Beispiel: "Normaler Text mit unterstrichener Passage und weiterem Text."
     → spans: [
         {"text": "Normaler Text mit ", "bold": false, "italic": false, "underline": false},
         {"text": "unterstrichener Passage", "bold": false, "italic": false, "underline": true},
         {"text": " und weiterem Text.", "bold": false, "italic": false, "underline": false}
       ]

6. **Platzhalter**: `FELD_*` wörtlich übernehmen. Auch bei OCR-Fehlern die
   offensichtlich gemeinte Form zurückgeben:
     "FELD__Name" → "FELD_Schuldner_Name"
     "FELD Akte LastGAW" → "FELD_Akte_LastGAW"
     "FELD_Akte_LastGAVV" → "FELD_Akte_LastGAW"
     "FELD_Schuldners-Schuldnerin" → "FELD_Schuldners_Schuldnerin"

7. **Listen**: Aufzählungen mit "-" oder "•" als eigene Absätze, jeder beginnend
   mit "- ".

8. **Weglassen**: Titel oben auf Seite 1 (z.B. "Bankenanfrage" als Überschrift)
   NICHT in den Output aufnehmen — der Dateiname reicht.

9. **Anlagen-Block**: "Anlagen:" als eigener Absatz (oft unterstrichen), dann pro
   Eintrag ("I. Gutachtenbeschluss", "II. Schweigepflichtentbindungserklärung")
   je ein Absatz mit Tab zwischen Nummer und Text: "I.\tGutachtenbeschluss".

Antworte NUR mit dem JSON, keine Einleitung, keine Code-Fences."""


def render_page_png(page: fitz.Page, dpi: int = 300) -> bytes:
    pix = page.get_pixmap(dpi=dpi)
    return pix.tobytes("png")


def call_claude(image_bytes: bytes, client: anthropic.Anthropic) -> dict[str, Any]:
    model = os.environ.get("EXTRACTION_MODEL", "claude-sonnet-4-5")
    # Strip the "-default" suffix used by Langdock; use bare model name for direct API
    model = model.removesuffix("-default")
    image_b64 = base64.standard_b64encode(image_bytes).decode()

    resp = client.messages.create(
        model=model,
        max_tokens=8192,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": image_b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": "Extrahiere diese Vorlagen-Seite als strukturiertes JSON gemäß Schema.",
                    },
                ],
            }
        ],
    )
    text = resp.content[0].text.strip()
    # Strip markdown code fences if present
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```\s*$", "", text)
    return json.loads(text)


def apply_post_fixes(text: str, extra_fixes: list[tuple[re.Pattern, str]] | None = None) -> str:
    for rx, repl in extra_fixes or []:
        text = rx.sub(repl, text)
    for rx, repl in POST_FIXES:
        text = rx.sub(repl, text)
    return text


def add_run_with_breaks(p, text: str, bold: bool, italic: bool, underline: bool) -> None:
    """Add a run that may contain \\n — each newline becomes a soft break."""
    parts = text.split("\n")
    for pi, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if underline:
                run.underline = True
        if pi < len(parts) - 1:
            br_run = p.add_run()
            br_run.add_break()


def write_docx(paragraphs: list[dict], out_path: Path,
               extra_fixes: list[tuple[re.Pattern, str]] | None = None,
               force_plain_prefixes: list[str] | None = None) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)

    force_plain_prefixes = force_plain_prefixes or []

    # Pre-compute paragraph texts for section detection
    texts: list[str] = []
    for para in paragraphs:
        texts.append("".join(s.get("text", "") for s in para.get("spans", [])))

    closing_start_idx: int | None = None
    anlagen_start_idx: int | None = None
    for i, t in enumerate(texts):
        stripped = t.strip()
        if closing_start_idx is None and "FELD_Bet_GrussBriefende" in t:
            closing_start_idx = i
        if anlagen_start_idx is None and stripped in {"Anlagen:", "Anlage:"}:
            anlagen_start_idx = i

    for i, para in enumerate(paragraphs):
        in_closing = closing_start_idx is not None and i >= closing_start_idx
        in_anlagen = anlagen_start_idx is not None and i >= anlagen_start_idx
        p = doc.add_paragraph()
        # End-Gruß/Signatur + Anlagen-Liste linksbündig, Fließtext Blocksatz
        p.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                       if (in_closing or in_anlagen)
                       else WD_ALIGN_PARAGRAPH.JUSTIFY)
        for span in para.get("spans", []):
            text = apply_post_fixes(span.get("text", ""), extra_fixes)
            if not text:
                continue
            bold = bool(span.get("bold"))
            italic = bool(span.get("italic"))
            underline = bool(span.get("underline"))
            if bold and any(text.lstrip().startswith(pfx) for pfx in force_plain_prefixes):
                bold = False
            add_run_with_breaks(p, text, bold, italic, underline)

        # Leerabsatz zwischen logischen Blöcken — NICHT innerhalb der Anlagen-Liste
        # (Anlagen: + I./II./III. → direkt untereinander, kein Leerabsatz)
        is_last = i >= len(paragraphs) - 1
        next_in_anlagen = (anlagen_start_idx is not None and (i + 1) > anlagen_start_idx)
        if is_last:
            continue
        if in_anlagen and next_in_anlagen:
            continue  # Skip blank paragraph inside Anlagen block
        empty = doc.add_paragraph()
        empty.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                           if in_closing or in_anlagen
                           else WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(str(out_path))


def convert_pdf(pdf_path: Path, docx_path: Path, client: anthropic.Anthropic,
                extra_fixes: list[tuple[re.Pattern, str]] | None = None,
                force_plain_prefixes: list[str] | None = None) -> None:
    pdf = fitz.open(str(pdf_path))
    all_paragraphs: list[dict] = []
    for page_i, page in enumerate(pdf):
        print(f"    Seite {page_i+1}/{len(pdf)} → Claude …", end="", flush=True)
        img = render_page_png(page)
        result = call_claude(img, client)
        paras = result.get("paragraphs", [])
        print(f" {len(paras)} Absätze")
        all_paragraphs.extend(paras)
    pdf.close()
    write_docx(all_paragraphs, docx_path, extra_fixes, force_plain_prefixes)


def extract_placeholders(docx_path: Path) -> tuple[dict[str, int], list[str]]:
    doc = Document(str(docx_path))
    TOK = re.compile(r"FELD_[A-Za-zÄÖÜäöüß0-9]+(?:_[A-Za-zÄÖÜäöüß0-9]+)*")
    text = "\n".join(p.text for p in doc.paragraphs)
    known: dict[str, int] = {}
    unknown: list[str] = []
    for m in TOK.findall(text):
        if m in KNOWN_FIELDS:
            known[m] = known.get(m, 0) + 1
        else:
            unknown.append(m)
    return known, unknown


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--only", help="Nur diese DOCX-Datei generieren (z.B. 'bausparkassen-anfrage.docx')")
    args = ap.parse_args()

    if not SRC.exists():
        print(f"Source dir missing: {SRC}", file=sys.stderr)
        return 1
    DST.mkdir(parents=True, exist_ok=True)

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    base_url = os.environ.get("ANTHROPIC_BASE_URL")
    if not api_key:
        print("ANTHROPIC_API_KEY fehlt in .env", file=sys.stderr)
        return 1
    # Für Claude Vision benutzen wir den direkten Anthropic-Endpoint, nicht Langdock
    # (Langdock-Proxy unterstützt Multimodal nicht zuverlässig).
    client = anthropic.Anthropic(api_key=api_key)

    summary: list[tuple[str, dict[str, int], list[str]]] = []
    for pdf_name, docx_name in PDF_MAP.items():
        if args.only and docx_name != args.only:
            continue
        if docx_name in SKIP_FILES and not args.only:
            print(f"  [skip] {docx_name} (manuell nachbearbeitet)")
            continue
        pdf_path = SRC / pdf_name
        docx_path = DST / docx_name
        if not pdf_path.exists():
            print(f"  ! MISSING: {pdf_name}", file=sys.stderr)
            continue
        print(f"  {pdf_name} → {docx_name}")
        try:
            convert_pdf(
                pdf_path, docx_path, client,
                PER_FILE_FIXES.get(docx_name),
                FORCE_PLAIN_PREFIXES.get(docx_name),
            )
        except Exception as e:
            print(f"    FEHLER: {e}", file=sys.stderr)
            continue
        known, unknown = extract_placeholders(docx_path)
        summary.append((docx_name, known, unknown))

    print()
    print("=" * 60)
    print("Ergebnis pro Template")
    print("=" * 60)
    for name, known, unknown in summary:
        print(f"\n{name}")
        for ph in sorted(known):
            c = known[ph]
            print(f"  ✓ {ph}" + (f"  ×{c}" if c > 1 else ""))
        if unknown:
            print("  Unbekannte Tokens (bitte prüfen):")
            for u in sorted(set(unknown)):
                print(f"    ? {u}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
