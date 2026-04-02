#!/usr/bin/env python3
"""
Create the golden Gutachten template for natürliche Person.

Takes the existing DOCX template (with correct formatting, headers, footers, styles)
and replaces paragraph/table content to match the real Geldt Gutachten structure.

Uses KI_* placeholders for data fields and [[SLOT_NNN: description]] for AI-filled narrative.
"""

import copy
import os
import sys
from pathlib import Path
from lxml import etree
from docx import Document

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "gutachtenvorlagen" / "Gutachten Muster natürliche Person.docx"
OUTPUT_PATH = TEMPLATE_PATH  # Overwrite in place

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def qn(tag):
    """Qualified name for w: namespace."""
    return f"{{{W_NS}}}{tag}"


def set_paragraph_text(para, text):
    """
    Replace ALL runs in a paragraph with a single run containing the new text.
    Preserves the paragraph properties (pPr) and the first run's formatting (rPr).
    """
    elem = para._element

    # Save pPr
    pPr = elem.find(qn("pPr"))

    # Collect first run's rPr for reuse
    first_rPr = None
    first_run = elem.find(qn("r"))
    if first_run is not None:
        rPr = first_run.find(qn("rPr"))
        if rPr is not None:
            first_rPr = copy.deepcopy(rPr)

    # Remove everything except pPr
    for child in list(elem):
        tag = child.tag if isinstance(child.tag, str) else ""
        if tag == qn("pPr"):
            continue
        elem.remove(child)

    # Create new run with text
    run = etree.SubElement(elem, qn("r"))
    if first_rPr is not None:
        run.insert(0, first_rPr)
    t = etree.SubElement(run, qn("t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def remove_paragraph(para):
    """Remove a paragraph element from the document body."""
    elem = para._element
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)


def insert_paragraph_after(doc, ref_para, text, style_name=None, copy_rPr_from=None):
    """
    Insert a new paragraph after ref_para with given text.
    If style_name is given, set the style.
    If copy_rPr_from is a paragraph, copy first run's rPr.
    Returns the new paragraph element.
    """
    body = doc.element.body
    ref_elem = ref_para._element

    # Create new w:p
    new_p = etree.Element(qn("p"))
    pPr = etree.SubElement(new_p, qn("pPr"))

    if style_name:
        pStyle = etree.SubElement(pPr, qn("pStyle"))
        pStyle.set(qn("val"), style_name)

    # Copy paragraph formatting from reference if same style
    if copy_rPr_from is not None:
        ref_pPr = copy_rPr_from._element.find(qn("pPr"))
        if ref_pPr is not None:
            # Copy indent, jc etc
            for child in ref_pPr:
                tag = child.tag if isinstance(child.tag, str) else ""
                if tag != qn("pStyle"):
                    pPr.append(copy.deepcopy(child))

    # Create run
    run = etree.SubElement(new_p, qn("r"))

    # Copy rPr from reference paragraph's first run
    if copy_rPr_from is not None:
        first_run = copy_rPr_from._element.find(qn("r"))
        if first_run is not None:
            rPr = first_run.find(qn("rPr"))
            if rPr is not None:
                run.insert(0, copy.deepcopy(rPr))

    t = etree.SubElement(run, qn("t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Insert after ref
    ref_elem.addnext(new_p)
    return new_p


def insert_empty_paragraph_after(doc, ref_para, style_name=None):
    """Insert an empty paragraph after ref_para."""
    body = doc.element.body
    ref_elem = ref_para._element
    new_p = etree.Element(qn("p"))
    if style_name:
        pPr = etree.SubElement(new_p, qn("pPr"))
        pStyle = etree.SubElement(pPr, qn("pStyle"))
        pStyle.set(qn("val"), style_name)
    ref_elem.addnext(new_p)
    return new_p


def set_table_cell_text(cell, text):
    """Set text of a table cell, preserving first paragraph formatting."""
    # Clear all paragraphs except first
    for para in cell.paragraphs[1:]:
        remove_paragraph(para)
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)


def get_para_by_index(doc, idx):
    """Get paragraph by index, with bounds check."""
    if 0 <= idx < len(doc.paragraphs):
        return doc.paragraphs[idx]
    return None


# ============================================================================
# MAIN TRANSFORMATION
# ============================================================================

def transform_template():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))

    paras = doc.paragraphs
    total = len(paras)
    print(f"Template has {total} paragraphs, {len(doc.tables)} tables")

    # ========================================================================
    # PAGE 1 — Cover page (paragraphs 0-29)
    # The cover page structure is already good. Fix:
    # - Para 4: "des/der" instead of "Artikel" for Vermögen reference
    # ========================================================================

    # Para 4: change KI_Schuldner_Artikel to KI_Schuldner_des_der
    set_paragraph_text(paras[4], "in dem Insolvenzantragsverfahren über das Vermögen KI_Schuldner_des_der")

    # Para 29: fix spacing in AZ reference
    set_paragraph_text(paras[29], "-KI_Akte_GerichtAZ-")

    # ========================================================================
    # PAGE 2 — Sachverständiger + Ergebnis (paragraphs 30-60)
    # Keep Sachverständiger block (30-39) as is — already correct.
    # Add Datum line after Aktenzeichen.
    # Replace result summary placeholder paragraphs (44-60) with proper slots.
    # ========================================================================

    # Add "Datum:" line — we need to insert after para 39 (Aktenzeichen)
    # First let's set para 41 (currently empty with runs) to Datum
    set_paragraph_text(paras[41], "                                          Datum: \t[[SLOT_001: Datum Gutachten]]")

    # Para 44: "welches zu dem Ergebnis kommt:" — keep as is

    # Replace the result placeholder paragraphs (46-60) with proper SLOT content
    # Para 46-52 are empty placeholders, 58="das Insolvenzverfahren", 60="zu eröffnen."
    # We'll reuse these for the Ergebnis slots

    set_paragraph_text(paras[46], "[[SLOT_002: Ergebnis Punkt 1 - Feststellung Zahlungsunfähigkeit/Überschuldung]]")
    set_paragraph_text(paras[48], "[[SLOT_003: Ergebnis Punkt 2 - Verfahrenskostendeckung]]")
    set_paragraph_text(paras[50], "[[SLOT_004: Ergebnis Punkt 3 - Empfehlung Sicherungsmaßnahmen]]")
    set_paragraph_text(paras[52], "[[SLOT_005: Ggf. Empfehlung Eröffnungsdatum und Verfahrenskostenstundung]]")

    # Remove the old "das Insolvenzverfahren" / "zu eröffnen." lines
    # Para 56, 58, 59, 60 — clear them
    for idx in [54, 55, 56, 57, 58, 59, 60]:
        if idx < total:
            set_paragraph_text(paras[idx], "")
            # Remove all runs to make truly empty
            elem = paras[idx]._element
            for child in list(elem):
                tag = child.tag if isinstance(child.tag, str) else ""
                if tag != qn("pPr"):
                    elem.remove(child)

    # ========================================================================
    # PAGE 3-4 — Inhaltsverzeichnis (paragraphs 61-119)
    # TOC is regenerated by Word. Leave as is.
    # ========================================================================

    # ========================================================================
    # PAGE 5 — A. Gutachtenauftrag / I. Auftrag (paragraphs 121-136)
    # Heading 1 "Gutachtenauftrag und Grundlagen für das Gutachten" — keep
    # Heading 2 "Auftrag" — keep
    # ========================================================================

    # Para 126: Auftrag text — update to match gold standard
    set_paragraph_text(paras[126],
        "Das Amtsgericht KI_Gericht_Ort hat KI_Verwalter_den_die Unterzeichner mit Beschluss vom "
        "KI_Akte_BeschlussDat beauftragt, in dem")

    # Para 128: "Insolvenzantragsverfahren..." — fix to use des_der
    set_paragraph_text(paras[128], "Insolvenzantragsverfahren über das Vermögen KI_Schuldner_des_der")

    # Para 130: KI_Akte_Bezeichnung — replace with Schuldner block
    set_paragraph_text(paras[130], "KI_Schuldner_NameVorname")

    # Use paragraphs 131-135 for the rest of the Schuldner block
    # 131 was empty, 132 was "ein schriftliches...", 133-135 empty
    # We need to insert more lines. Let's reuse what we have.

    # Actually the structure should be:
    # KI_Schuldner_NameVorname (130)
    # geb. KI_Schuldner_Geburtsdatum (131 — currently empty)
    # KI_Schuldner_Adr (132 — currently "ein schriftliches...")
    # [empty] (133)
    # Inh. KI_Schuldner_Firma (134)
    # KI_Schuldner_Betriebsstaette (135)

    # But we also need the questions a) b) c) and the closing line.
    # Let's use a different approach: modify in sequence.

    set_paragraph_text(paras[131], "geb. KI_Schuldner_Geburtsdatum")
    set_paragraph_text(paras[132], "KI_Schuldner_Adr")
    set_paragraph_text(paras[133], "")
    set_paragraph_text(paras[134], "Inh. KI_Schuldner_der_die KI_Schuldner_Firma")
    set_paragraph_text(paras[135], "KI_Schuldner_Betriebsstaette")

    # Now we need to add more content after para 135: the questions block
    # We have para 136 which is "In Erfüllung dieses Auftrages..."
    # We need to insert between 135 and 136:
    # - empty line
    # - "- nachfolgend „Antragsteller" -"
    # - empty line
    # - "ein schriftliches Gutachten..."
    # - empty line
    # - question a)
    # - empty line
    # - "Falls ja."
    # - empty line
    # - question b)
    # - empty line
    # - question c)
    # - empty line

    # We'll insert paragraphs after 135, building backwards
    # Actually, let's insert forward by chaining addnext
    ref = paras[135]._element

    lines_to_insert = [
        "",
        '- nachfolgend \u201eAntragsteller\u201c -',
        "",
        "ein schriftliches Gutachten zur Beantwortung folgender Fragen zu erstellen:",
        "",
        "a) Liegen Tatsachen vor, wonach der Schluss auf (drohende) Zahlungsunfähigkeit des Antragstellers gerechtfertigt ist?",
        "",
        "Falls ja.",
        "",
        "b) Ist eine die Verfahrenskosten (§ 54 InsO) deckende Masse vorhanden? Dabei sind auch insolvenzspezifische Ansprüche (Haftung von verantwortlichen Organen, Anfechtungsansprüche) zu prüfen und darzustellen, wann die Zahlungsunfähigkeit / Überschuldung eingetreten ist.",
        "",
        "c) Erscheinen vorläufige Anordnungen zur Sicherung der Masse (allgemeines Veräußerungsverbot, vorläufige Verwaltung, Postsperre usw.) erforderlich?",
        "",
    ]

    # Get pPr from para 132 (was the questions paragraph — has indent)
    ref_pPr = paras[132]._element.find(qn("pPr"))

    for line_text in lines_to_insert:
        new_p = etree.Element(qn("p"))
        if ref_pPr is not None:
            new_p.append(copy.deepcopy(ref_pPr))
        if line_text:
            run = etree.SubElement(new_p, qn("r"))
            t = etree.SubElement(run, qn("t"))
            t.text = line_text
            if line_text[0] == " " or line_text[-1] == " ":
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        ref.addnext(new_p)
        ref = new_p

    # Para 136 already says "In Erfüllung dieses Auftrages..." — update it
    set_paragraph_text(paras[136],
        "In Erfüllung dieses Auftrages hat KI_Verwalter_der_die Unterzeichner das nachfolgende "
        "Gutachten erstellt und erstattet hiermit Bericht über den Verlauf des Antragsverfahrens:")

    # ========================================================================
    # PAGE 6 — II. Grundlagen (paragraphs 137-165)
    # ========================================================================

    # Para 137: Heading 2 "Grundlagen" — keep

    # Para 139: Grundlagen intro — update
    set_paragraph_text(paras[139],
        "Dem Beschluss ging ein [[SLOT_006: Eigenantrag/Fremdantrag]] KI_Schuldner_des_der "
        "Antragstellers vom KI_Akte_VorlVerfahrDat voraus, welcher am "
        "[[SLOT_007: Eingangsdatum bei Gericht]] bei dem Amtsgericht KI_Gericht_Ort einging.")

    # Insert second paragraph of Grundlagen after 139
    ref = paras[139]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)  # empty line
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = ("Dieses Gutachten basiert auf der Auswertung der Gerichtsakten, Gesprächen mit den unten "
              "genannten Auskunftspersonen, Inaugenscheinnahmen, den Ergebnissen des beauftragten "
              "Sachverständigen und den vom Antragsteller vorgelegten Unterlagen.")
    ref.addnext(new_p)

    # Para 141: Heading 3 "Unterlagen" — keep

    # Replace Unterlagen content (paras 143-149)
    # Para 143: was yellow-highlighted empty → replace with Unterlagen content
    set_paragraph_text(paras[143],
        "Das Insolvenzgericht hat KI_Verwalter_dem_der Unterzeichner am "
        "[[SLOT_008: Datum Aktenübersendung]] die Gerichtsakte übersandt.")

    # Para 145: was yellow multi-line → first meeting
    set_paragraph_text(paras[145],
        "[[SLOT_009: Beschreibung des ersten Besprechungstermins mit dem Schuldner]]")

    # Para 147: was "Datensicherung" text → replace with unterlagen list intro
    set_paragraph_text(paras[147],
        "Dem Unterzeichner standen darüber hinaus folgende Unterlagen zur Einsichtnahme zur Verfügung:")

    # Para 149: was Steuerberater text → replace with unterlagen list
    set_paragraph_text(paras[149],
        "[[SLOT_010: Liste der Unterlagen (Buchungskonten, OPOS-Listen, Grundbuchauszug, Creditreform etc.)]]")

    # Para 151: Heading 3 "Auskunftspersonen" — keep

    # Para 153: was yellow-highlighted → replace
    set_paragraph_text(paras[153],
        "Darüber hinaus standen KI_Verwalter_dem_der Unterzeichner folgende Personen für Auskünfte zur Verfügung:")

    # Insert Auskunftspersonen list slot after 153
    ref = paras[153]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = "[[SLOT_011: Liste der Auskunftspersonen mit Name und Organisation]]"
    ref.addnext(new_p)

    # Para 155: Heading 3 "VID Erklärung" — keep
    # Paras 157-163: VID boilerplate — KEEP AS IS (already correct)

    # ========================================================================
    # PAGE 8-9 — B. Allgemeine Angaben / Tables
    # The tables are already structured. Update cell contents.
    # ========================================================================

    # Table 0: Statistische Angaben (15 rows x 2 cols)
    t0 = doc.tables[0]

    # Row 0: Name → KI_Schuldner_NameVorname (was KI_Schuldner_Name)
    set_table_cell_text(t0.rows[0].cells[1], "KI_Schuldner_NameVorname")

    # Row 1: Handelsfirma → keep KI_Schuldner_Firma
    # Row 2: Anschrift → Add both addresses
    set_table_cell_text(t0.rows[2].cells[0], "Anschrift (privat)")
    # Row 2 value stays KI_Schuldner_Adr

    # We need to add a row for Betriebsstaette — but modifying table structure is complex.
    # Instead, let's update existing rows with the right content.

    # Row 3: Handelsregister → keep KI_Schuldner_HRB, but add slot
    set_table_cell_text(t0.rows[3].cells[1], "[[SLOT_012: Handelsregister ja/nein mit Details]]")

    # Row 4: Insolvenzforderungen → update with all three values
    set_table_cell_text(t0.rows[4].cells[1],
        "KI_Forderungen_Gesamt / KI_Forderungen_Gesichert / [[SLOT_013: Nachrangig]]")

    # Row 5: Antragsteller → add slot for details
    set_table_cell_text(t0.rows[5].cells[1],
        "KI_Antragsteller_Name [[SLOT_014: Eigenantrag/Fremdantrag Details]]")

    # Row 6: Antragsgrund → keep KI_Verfahren_Eroeffnungsgrund
    # Row 7: Internationaler Bezug → add slot
    set_table_cell_text(t0.rows[7].cells[1], "[[SLOT_015: Internationaler Bezug]]")
    # Row 8: Eigenverwaltung → add slot
    set_table_cell_text(t0.rows[8].cells[1], "[[SLOT_016: Eigenverwaltung]]")
    # Row 9: Geschäftszweig → simplify
    set_table_cell_text(t0.rows[9].cells[1], "[[SLOT_017: Geschäftszweig und Unternehmensgegenstand]]")
    # Row 10: Unternehmensgegenstand → keep KI_Schuldner_Firma or make slot
    set_table_cell_text(t0.rows[10].cells[1], "[[SLOT_018: Unternehmensgegenstand Beschreibung]]")
    # Row 11: Arbeitnehmer → keep KI_Arbeitnehmer_Anzahl
    # Row 12: Betriebsrat → add slot
    set_table_cell_text(t0.rows[12].cells[1], "[[SLOT_019: Betriebsrat ja/nein]]")
    # Row 13: Restschuldbefreiung → add slot
    set_table_cell_text(t0.rows[13].cells[1], "[[SLOT_020: Restschuldbefreiung ja/nein]]")
    # Row 14: Verfahrenskostenstundung → add slot
    set_table_cell_text(t0.rows[14].cells[1], "[[SLOT_021: Verfahrenskostenstundung ja/nein]]")

    # Table 1: Steuerrechtliche Angaben (5 rows x 2 cols)
    t1 = doc.tables[1]
    # Row 0: Finanzamt → keep KI_Finanzamt
    # Row 1: Steuer-Nr. → keep KI_Steuernummer, add USt-ID
    set_table_cell_text(t1.rows[1].cells[0], "Steuer-Nr. / USt-ID-Nr.")
    set_table_cell_text(t1.rows[1].cells[1], "KI_Steuernummer / [[SLOT_022: USt-ID]]")
    # Row 2: Wirtschaftsjahr / USt → add slots
    set_table_cell_text(t1.rows[2].cells[1], "[[SLOT_023: Wirtschaftsjahr]] / [[SLOT_024: USt-Versteuerung]]")
    # Row 3: Steuerliche Organschaft → slot
    set_table_cell_text(t1.rows[3].cells[1], "[[SLOT_025: Steuerliche Organschaft]]")
    # Row 4: Letzter Jahresabschluss → keep KI_Letzter_Jahresabschluss

    # Table 2: Sonstige Angaben (18 rows x 2 cols)
    t2 = doc.tables[2]
    # Row 0: Geburtsdatum/Geburtsort → add Geburtsort slot
    set_table_cell_text(t2.rows[0].cells[1], "KI_Schuldner_Geburtsdaten / [[SLOT_026: Geburtsort]]")
    # Row 1: Konfession → slot
    set_table_cell_text(t2.rows[1].cells[1], "[[SLOT_027: Konfession]]")
    # Row 2: Ausbildung → keep KI_Ausbildung
    # Row 3: Unterhaltspflichten → slot
    set_table_cell_text(t2.rows[3].cells[1], "[[SLOT_028: Kinder und Unterhalt Details mit Geburtsdaten]]")
    # Row 4: Ehegatten → slot
    set_table_cell_text(t2.rows[4].cells[1], "[[SLOT_029: Familienstand Details mit Geburtsdatum]]")
    # Row 5: Telefon → keep KI_Schuldner_Telefon
    # Row 6: Mobiltelefon → keep
    # Row 7: E-Mail → keep KI_Schuldner_Email
    # Row 8: SVTraeger → keep
    # Row 9: Betriebsnummer → keep
    # Row 10: Berufsgenossenschaft → slot
    set_table_cell_text(t2.rows[10].cells[1], "[[SLOT_030: Berufsgenossenschaft und Mitgliedsnummer]]")
    # Row 11: älteste Forderung → keep KI_Aelteste_Forderung
    # Row 12: erste bilanzielle Überschuldung → slot
    set_table_cell_text(t2.rows[12].cells[1], "[[SLOT_031: Erste bilanzielle Überschuldung]]")
    # Row 13: Steuerberater → keep
    # Row 14: wirtschaftlich Berechtigter → slot
    set_table_cell_text(t2.rows[14].cells[1], "[[SLOT_032: Wirtschaftlich Berechtigter nach Transparenzregister]]")
    # Row 15: Bankverbindungen → keep KI_Bankverbindungen
    # Row 16: Insolvenzsonderkonto → update with both fields
    set_table_cell_text(t2.rows[16].cells[1], "KI_ANDERKONTO_IBAN / KI_Anderkonto_Bank")
    # Row 17: Gerichtsvollzieher → keep KI_Gerichtsvollzieher

    # ========================================================================
    # B.II. Angaben zum Verfahren (para 180+)
    # ========================================================================

    # Para 182: was yellow-highlighted empty → Verfahrenshistorie
    set_paragraph_text(paras[182],
        "[[SLOT_033: Verfahrenshistorie - Antragsdatum, Eingangsdatum, Beschlussdatum, "
        "Bestellung, ggf. Restschuldbefreiung, ggf. Verfahrenskostenstundung]]")

    # Para 184: Heading 3 "Internationale Zuständigkeit" — keep
    # Para 186: International jurisdiction text — keep (boilerplate)
    # Para 188: EuInsVO explanation — keep (boilerplate)
    # Para 190: Art. 3 Abs. 1 EuInsVO — keep (boilerplate)
    # Para 192: Natürliche Person presumption — keep (boilerplate)

    # Para 196: Heading 3 "Örtliche Zuständigkeit..." — keep
    # Para 198: Örtliche Zuständigkeit boilerplate — keep

    # Para 200: was yellow → concrete jurisdiction reasoning
    set_paragraph_text(paras[200],
        "[[SLOT_034: Konkrete Zuständigkeitsbegründung - wo ist der Geschäftsbetrieb, Gerichtsbezirk]]")

    # Para 202: was yellow → conclusion
    set_paragraph_text(paras[202],
        "Das Amtsgericht KI_Gericht_Ort ist daher örtlich zuständig.")

    # Para 204: Heading 3 "Sicherungsmaßnahmen" — keep
    # Para 206: was yellow → Sicherungsmaßnahmen content
    set_paragraph_text(paras[206],
        "[[SLOT_035: Sicherungsmaßnahmen - Beschluss, Datum, Maßnahmen "
        "(Zustimmungsvorbehalt, Zwangsvollstreckungsverbot, Kontensperre etc.)]]")

    # Para 208: Heading 3 "Einzelermächtigungen" — keep
    # Para 210: was yellow → Einzelermächtigungen content
    set_paragraph_text(paras[210],
        "[[SLOT_036: Einzelermächtigungen des vorläufigen Insolvenzverwalters]]")

    # ========================================================================
    # C. Informationen über das Unternehmen (para 213+)
    # ========================================================================

    # Para 213: Heading 1 "Informationen über das Unternehmen" — keep

    # C.I. Unternehmensgegenstand (para 216)
    # Para 216: Heading 2 — update to "Persönliche Verhältnisse und Unternehmensgegenstand"
    set_paragraph_text(paras[216], "Persönliche Verhältnisse und Unternehmensgegenstand")

    # Para 218: was slot → narrative
    set_paragraph_text(paras[218],
        "[[SLOT_037: Persönliche Verhältnisse des Schuldners und Unternehmensgegenstand - "
        "Biographie, Geschäftstätigkeit, Entwicklung des Unternehmens]]")

    # C.II. Finanzierungsstruktur (para 220)
    # Para 222: slot → narrative
    set_paragraph_text(paras[222],
        "[[SLOT_038: Finanzierungsstruktur - Bankverbindungen, Kreditlinien, "
        "Darlehen mit Konditionen und Sicherheiten pro Bank]]")

    # C.III. Wesentliche Vertragsverhältnisse (para 224)
    # 1. Arbeitsrechtliche Verhältnisse (para 226)
    set_paragraph_text(paras[228],
        "[[SLOT_039: Arbeitsrechtliche Verhältnisse - Anzahl Arbeitnehmer, "
        "Auszubildende, Betriebsrat, Kündigungen, Sozialplan]]")
    # Clear the old SLOT_102
    set_paragraph_text(paras[230], "")

    # 2. Gewerberaummietverträge (para 232)
    set_paragraph_text(paras[234],
        "[[SLOT_040: Gewerberaummietverträge / Pachtverträge - Objekt, Vermieter, "
        "Miete, Laufzeit, Kündigungsstatus]]")

    # 3. Versicherungsverträge (para 236)
    set_paragraph_text(paras[238],
        "[[SLOT_041: Versicherungsverträge - Art, Versicherer, Beiträge, Status]]")

    # 4. Versorgungsverträge (para 240)
    set_paragraph_text(paras[242],
        "[[SLOT_042: Versorgungsverträge - Strom, Gas, Wasser, Telekommunikation]]")

    # 5. Leasingverträge (para 244)
    set_paragraph_text(paras[246],
        "[[SLOT_043: Leasingverträge / Mietverträge über bewegliche Gegenstände]]")

    # 6. Factoring (para 248)
    set_paragraph_text(paras[250],
        "[[SLOT_044: Factoring-Vereinbarungen oder 'Factoring-Vereinbarungen bestehen nicht.']]")

    # 7. Sonstige Verträge (para 252)
    set_paragraph_text(paras[254],
        "[[SLOT_045: Sonstige Verträge oder 'Sonstige wesentliche Vertragsverhältnisse sind nicht bekannt.']]")

    # C.IV. Arbeitsrechtliche Verhältnisse [wenn größerer Betrieb] (para 256)
    # This is a duplicate heading for larger companies — replace content
    set_paragraph_text(paras[258],
        "[[SLOT_046: Erweiterte arbeitsrechtliche Verhältnisse bei größerem Betrieb - "
        "kollektivarbeitsrechtliche und betriebsverfassungsrechtliche Aspekte]]")
    set_paragraph_text(paras[260], "")

    # C.V. Zahlungsverkehr (para 262)
    set_paragraph_text(paras[264],
        "[[SLOT_047: Zahlungsverkehr - Kontenübersicht mit Kreditinstitut, Kontonummer, "
        "Saldo bei Antragstellung, Kontensperrung]]")

    # Clear old paras 266-269
    for idx in [266, 267, 268]:
        set_paragraph_text(paras[idx], "")
    set_paragraph_text(paras[269], "")

    # Table 3 (Konten) — update with generic slots
    t3 = doc.tables[3]
    # Header row
    set_table_cell_text(t3.rows[0].cells[0], "Konto")
    set_table_cell_text(t3.rows[0].cells[1], "Stand bei Antragstellung")
    # Data rows
    set_table_cell_text(t3.rows[1].cells[0], "[[SLOT_048: Kreditinstitut 1 mit Kontonummer]]")
    set_table_cell_text(t3.rows[1].cells[1], "[[SLOT_049: Saldo 1]] EUR")
    set_table_cell_text(t3.rows[2].cells[0], "[[SLOT_050: Kreditinstitut 2 mit Kontonummer]]")
    set_table_cell_text(t3.rows[2].cells[1], "[[SLOT_051: Saldo 2]] EUR")
    set_table_cell_text(t3.rows[3].cells[0], "[[SLOT_052: Weitere Konten]]")
    set_table_cell_text(t3.rows[3].cells[1], "[[SLOT_053: Saldo weitere]] EUR")

    # C.VI. Rechnungswesen (para 272)
    # 1. Buchhaltung vor Antragstellung (para 274)
    set_paragraph_text(paras[276],
        "[[SLOT_054: Buchhaltung vor Antragstellung - wer hat geführt, Qualität, Vollständigkeit]]")
    set_paragraph_text(paras[278], "")  # Clear old Lohnbuchhaltung text

    # 2. Buchhaltung nach Antragstellung (para 280)
    set_paragraph_text(paras[282],
        "[[SLOT_055: Buchhaltung nach Antragstellung - Fortführung durch wen, Maßnahmen]]")

    # C.VII. Anhängige Rechtsstreitigkeiten (para 284)
    set_paragraph_text(paras[286],
        "[[SLOT_056: Anhängige Rechtsstreitigkeiten oder "
        "'Anhängige Rechtsstreitigkeiten sind dem Unterzeichner nicht bekannt.']]")

    # C.VIII. Betriebswirtschaftliche Verhältnisse (para 288)
    # 1. Wirtschaftliche Entwicklung und Krisenursache (para 290)
    set_paragraph_text(paras[292],
        "[[SLOT_057: Wirtschaftliche Entwicklung und Krisenursache - "
        "ausführliche Darstellung der Unternehmensentwicklung, Umsatzentwicklung, "
        "Krisenursachen, ggf. mit Jahresabschlusszahlen]]")

    # 2. Maßnahmen im Insolvenzeröffnungsverfahren (para 294)
    # a) Vermögenssicherung (para 296)
    set_paragraph_text(paras[298],
        "[[SLOT_058: Vermögenssicherung - Maßnahmen zur Sicherung des Schuldnervermögens]]")

    # b) Sicherung der Betriebsfortführung (para 300)
    set_paragraph_text(paras[302],
        "[[SLOT_059: Sicherung der Betriebsfortführung - Betriebsfortführung ja/nein, "
        "Maßnahmen, Personal, Aufträge, Lieferanten]]")

    # Clear the old detailed paras 304-321
    for idx in range(304, 322):
        if idx < total:
            set_paragraph_text(paras[idx], "")

    # c) Datensicherung → repurpose para 312 (List Paragraph "Datensicherung")
    # Already cleared above. Insert a slot for Datensicherung
    # Para 312 is List Paragraph style — set it
    set_paragraph_text(paras[312], "Datensicherung")
    set_paragraph_text(paras[314],
        "[[SLOT_060: Datensicherung - Umfang, Medium, Aufbewahrung]]")

    # 3. Sanierungsaussichten (para 323)
    set_paragraph_text(paras[325],
        "[[SLOT_061: Sanierungsaussichten - Analyse der Sanierungsfähigkeit, "
        "ggf. Fortführungsprognose, ggf. Verwertungskonzept]]")

    # ========================================================================
    # D. Vermögensverhältnisse (para 328+)
    # ========================================================================

    # D.I. Bewertungsmaßstab (para 331) — boilerplate, keep as is

    # D.II. Sachverständige Bewertung (para 335)
    set_paragraph_text(paras[337],
        "KI_Verwalter_Der_Die_Groß Unterzeichner hat die Bewertung und Inventarisierung "
        "des beweglichen Sachanlagevermögens in Auftrag gegeben.")

    # Para 339: was yellow → details
    set_paragraph_text(paras[339],
        "Die Aufnahme erfolgte zum Bewertungsstichtag des [[SLOT_062: Bewertungsstichtag]] "
        "durch den Sachverständigen [[SLOT_063: Sachverständiger Name und Adresse]].")

    # Insert reference to Anlage
    ref = paras[339]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = ("Inhaltlich nimmt KI_Verwalter_der_die Unterzeichner insoweit Bezug auf das "
              "als Anlage II beigefügte Gutachten vom [[SLOT_064: Datum Wertgutachten]].")
    ref.addnext(new_p)

    # D.III. Aktiva (para 341+)
    # 1. Anlagevermögen (para 343)
    # a) Immaterielles Vermögen (para 345)
    # aa) Goodwill (346) — keep label
    set_paragraph_text(paras[348],
        "[[SLOT_065: Goodwill / Firmenwert Beschreibung und Bewertung]]")
    # bb) selbst geschaffene Schutzrechte (350) — keep label
    set_paragraph_text(paras[352],
        "[[SLOT_066: Selbst geschaffene gewerbliche Schutzrechte Beschreibung/Wert]]")
    # cc) entgeltlich erworbene (354) — keep label
    set_paragraph_text(paras[356],
        "[[SLOT_067: Entgeltlich erworbene Konzessionen und Schutzrechte Beschreibung/Wert]]")

    # b) Sachanlagen (para 358)
    # aa) Grundstücke (359) — keep label
    set_paragraph_text(paras[361],
        "[[SLOT_068: Grundstücke und grundstücksgleiche Rechte - Beschreibung, "
        "Grundbuch, Verkehrswert, Belastungen]]")
    # bb) technische Anlagen (363) — keep label
    set_paragraph_text(paras[365],
        "[[SLOT_069: Technische Anlagen und Maschinen - Beschreibung, Bewertung]]")
    # cc) BGA (367) — keep label
    set_paragraph_text(paras[369],
        "[[SLOT_070: Betriebs- und Geschäftsausstattung - Beschreibung, Bewertung]]")

    # c) Finanzanlagen (para 371)
    # aa) Anteile verbundene Unternehmen (373)
    set_paragraph_text(paras[375],
        "[[SLOT_071: Anteile an verbundenen Unternehmen Beschreibung/Wert]]")
    # bb) Beteiligungen (377)
    set_paragraph_text(paras[379],
        "[[SLOT_072: Beteiligungen Beschreibung/Wert]]")
    # cc) Wertpapiere Anlagevermögen (381)
    set_paragraph_text(paras[383],
        "[[SLOT_073: Wertpapiere des Anlagevermögens Beschreibung/Wert]]")

    # 2. Umlaufvermögen (para 385)
    # a) Vorräte (para 387)
    # aa) Roh-/Hilfs-/Betriebsstoffe (389)
    set_paragraph_text(paras[391],
        "[[SLOT_074: Roh-, Hilfs- und Betriebsstoffe Beschreibung/Wert]]")
    # bb) unfertige Erzeugnisse (393)
    set_paragraph_text(paras[395],
        "[[SLOT_075: Unfertige Erzeugnisse, unfertige Leistungen Beschreibung/Wert]]")
    # cc) fertige Erzeugnisse (397)
    set_paragraph_text(paras[399],
        "[[SLOT_076: Fertige Erzeugnisse / Waren Beschreibung/Wert]]")

    # b) Forderungen und sonstige Vermögensgegenstände (para 401)
    # aa) and bb) — paras 402-409
    set_paragraph_text(paras[402],
        "[[SLOT_077: Forderungen aus Lieferungen und Leistungen - Bestand, Einbringlichkeit, Bewertung]]")
    set_paragraph_text(paras[404], "")
    set_paragraph_text(paras[406],
        "[[SLOT_078: Forderungen gegen verbundene Unternehmen]]")
    set_paragraph_text(paras[408], "")

    # cc) Forderungen Beteiligungsunternehmen (410)
    set_paragraph_text(paras[412],
        "[[SLOT_079: Forderungen gegen Beteiligungsunternehmen]]")

    # dd) sonstige Vermögensgegenstände (414)
    set_paragraph_text(paras[416],
        "[[SLOT_080: Sonstige Vermögensgegenstände Beschreibung/Wert]]")

    # c) Wertpapiere (para 418)
    set_paragraph_text(paras[419],
        "[[SLOT_081: Wertpapiere des Umlaufvermögens Beschreibung/Wert]]")

    # d) Liquide Mittel (para 421)
    set_paragraph_text(paras[422],
        "[[SLOT_082: Kassenbestand]]")
    set_paragraph_text(paras[424], "")
    set_paragraph_text(paras[426],
        "[[SLOT_083: Bankguthaben bei Antragstellung mit Details pro Konto]]")

    # Anderkonto text (428) — keep reference to KI_* fields
    set_paragraph_text(paras[428],
        "Der Unterzeichner hat bei der KI_Anderkonto_Bank ein Insolvenz-Sonderkonto "
        "mit der IBAN: KI_ANDERKONTO_IBAN eingerichtet. "
        "Das aktuelle Guthaben beträgt [[SLOT_084: Guthaben Anderkonto]] EUR.")

    # Additional accounts (430)
    set_paragraph_text(paras[430],
        "[[SLOT_085: Weitere Konten (Festgeld, Rückstellungskonto) mit IBAN und Guthaben]]")

    # 3. Zusammenfassung Aktiva (para 432)
    # a) Liquidationswerte (434)
    set_paragraph_text(paras[435],
        "[[SLOT_086: Zusammenfassung Liquidationswerte]]")
    # b) Fortführungswerte (437)
    set_paragraph_text(paras[438],
        "[[SLOT_087: Zusammenfassung Fortführungswerte]]")
    # c) Aussonderung (440)
    set_paragraph_text(paras[442],
        "[[SLOT_088: Aussonderungsrechte Beschreibung]]")
    # Table placeholder (444)
    set_paragraph_text(paras[444],
        "[[SLOT_089: Tabelle Aussonderung]]")
    # d) Absonderung (446)
    set_paragraph_text(paras[448],
        "[[SLOT_090: Absonderungsrechte Beschreibung]]")
    # Table placeholder (450)
    set_paragraph_text(paras[450],
        "[[SLOT_091: Tabelle Absonderung]]")

    # Tables 4 & 5 (Aktiva summary tables) — leave structure, they'll be filled dynamically

    # D.IV. Passiva (para 452)
    set_paragraph_text(paras[454],
        "[[SLOT_092: Passiva-Tabelle mit Forderungsübersicht nach Kategorien]]")

    # D.V. Drittsicherheiten (para 456)
    set_paragraph_text(paras[458],
        "[[SLOT_093: Drittsicherheiten - Bürgschaften, Patronatserklärungen, "
        "sonstige Sicherheiten Dritter]]")

    # ========================================================================
    # E. Eröffnungsgründe (para 461+)
    # ========================================================================

    # E.I. Zahlungsunfähigkeit (para 464)
    # 1. Definition (para 466) — KEEP ALL BOILERPLATE (paras 468, 470, 472)
    # 2. Vorliegen im Gutachtenfall (para 474)
    set_paragraph_text(paras[476],
        "[[SLOT_094: Vorliegen der Zahlungsunfähigkeit im konkreten Fall - "
        "Finanzstatus, Liquiditätslücke, fällige Verbindlichkeiten vs. verfügbare Mittel]]")

    # E.II. Überschuldung (para 478) — keep boilerplate (para 480) for natürliche Person

    # E.III. Eintritt der Insolvenzreife (para 482)
    # Para 484: was yellow → slot
    set_paragraph_text(paras[484],
        "[[SLOT_095: Eintritt der Insolvenzreife - Zeitpunkt und Begründung]]")
    # Para 486: keep the "abschließende Einschätzung" text as is

    # ========================================================================
    # F. Insolvenzspezifische Ansprüche (para 489+)
    # ========================================================================

    # F.I. Ansprüche aus Insolvenzanfechtung (para 492)
    set_paragraph_text(paras[494],
        "[[SLOT_096: Anfechtungsansprüche §§ 129 ff. InsO - "
        "identifizierte anfechtbare Rechtshandlungen, Lookback-Zeiträume, "
        "Erfolgsaussichten und geschätztes Volumen]]")

    # F.II. Kostenbeiträge gem. § 171 InsO (para 497)
    set_paragraph_text(paras[499],
        "[[SLOT_097: Kostenbeiträge gem. § 171 InsO - "
        "Sicherungsrechte und daraus resultierende Kostenbeiträge]]")

    # F.III. Zusammenfassung (para 501)
    set_paragraph_text(paras[503],
        "[[SLOT_098: Zusammenfassung insolvenzspezifische Ansprüche - "
        "Gesamtüberblick Anfechtung + Kostenbeiträge]]")

    # ========================================================================
    # G. Verfahrenskostendeckung (para 506+)
    # ========================================================================

    set_paragraph_text(paras[509],
        "[[SLOT_099: Verfahrenskostendeckung - Berechnung nach InsVV, "
        "Vergütung vorläufiges Verfahren, Vergütung eröffnetes Verfahren, "
        "Gerichtskosten, Gegenüberstellung mit verfügbarer Masse]]")

    # Table 6 (Verfahrenskosten) — update
    t6 = doc.tables[6]
    set_table_cell_text(t6.rows[0].cells[1], "[[SLOT_100: Vergütung vorläufiges Verfahren]] EUR")
    set_table_cell_text(t6.rows[1].cells[1], "[[SLOT_101: Vergütung eröffnetes Verfahren]] EUR")
    set_table_cell_text(t6.rows[2].cells[1], "[[SLOT_102: Gerichtskosten]] EUR")
    set_table_cell_text(t6.rows[3].cells[1], "KI_Verfahrenskosten_Gesamt EUR")

    # Clear remaining paras in G section
    for idx in [510, 511, 512, 513, 514, 515, 516]:
        if idx < total:
            set_paragraph_text(paras[idx], "")

    # ========================================================================
    # H. Ergebnis und Empfehlung (para 517+)
    # ========================================================================

    set_paragraph_text(paras[520],
        "[[SLOT_103: Ergebnis und Empfehlung - Nummerierte Feststellungen "
        "(1. Zahlungsunfähigkeit, 2. Kostendeckung, 3. Sicherungsmaßnahmen) "
        "und abschließende Empfehlung zur Eröffnung]]")

    # Signature block (para 524+) — keep "Mit freundlichen Grüßen,"
    # Update the signature
    set_paragraph_text(paras[526], "KI_Verwalter_Name")
    set_paragraph_text(paras[527], "als Sachverständiger und vorläufiger Insolvenzverwalter")
    set_paragraph_text(paras[528], "")  # Clear female variant (handled by comment rules)

    # Anlagenverzeichnis (para 531) — keep heading
    # Table 7 (Anlagen) — update
    t7 = doc.tables[7]
    set_table_cell_text(t7.rows[0].cells[1], "[[SLOT_104: Anlage I Bezeichnung]]")
    set_table_cell_text(t7.rows[1].cells[1], "[[SLOT_105: Anlage II Bezeichnung]]")
    set_table_cell_text(t7.rows[2].cells[1], "[[SLOT_106: Anlage III Bezeichnung]]")
    set_table_cell_text(t7.rows[3].cells[1], "[[SLOT_107: Anlage IV Bezeichnung]]")
    set_table_cell_text(t7.rows[4].cells[1], "[[SLOT_108: Anlage V Bezeichnung]]")

    # ========================================================================
    # SAVE
    # ========================================================================

    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(str(OUTPUT_PATH))
    print("Done!")

    # Verification: reload and check
    print("\n=== VERIFICATION ===")
    doc2 = Document(str(OUTPUT_PATH))
    print(f"Paragraphs: {len(doc2.paragraphs)}")
    print(f"Tables: {len(doc2.tables)}")

    # Count KI_* and SLOT markers
    ki_count = 0
    slot_count = 0
    for p in doc2.paragraphs:
        text = p.text
        ki_count += text.count("KI_")
        slot_count += text.count("[[SLOT_")
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                ki_count += text.count("KI_")
                slot_count += text.count("[[SLOT_")

    print(f"KI_* placeholders found: {ki_count}")
    print(f"[[SLOT_*]] markers found: {slot_count}")

    # Verify key headings still exist
    headings = [(p.style.name, p.text) for p in doc2.paragraphs
                if p.style and p.style.name.startswith("Heading")]
    print(f"Headings preserved: {len(headings)}")
    for style, text in headings:
        print(f"  {style}: {text[:80]}")


if __name__ == "__main__":
    transform_template()
