#!/usr/bin/env python3
"""
Create the golden Gutachten template for juristische Person.

Takes the existing DOCX template (with correct formatting, headers, footers, styles)
and replaces paragraph/table content to match the real freiraum 3 GmbH Gutachten structure.

Uses KI_* placeholders for data fields and [[SLOT_NNN: description]] for AI-filled narrative.
"""

import copy
import os
import sys
from pathlib import Path
from lxml import etree
from docx import Document

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PROJECT_ROOT / "gutachtenvorlagen" / "Gutachten Muster juristische Person.docx"
OUTPUT_PATH = TEMPLATE_PATH  # Overwrite in place

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def qn(tag):
    """Qualified name for w: namespace."""
    return f"{{{W_NS}}}{tag}"


def set_paragraph_text(para, text):
    """
    Replace ALL runs in a paragraph with a single run containing the new text.
    Preserves the paragraph properties (pPr) and the first run's formatting (rPr).
    """
    elem = para._element

    # Save pPr
    pPr = elem.find(qn("pPr"))

    # Collect first run's rPr for reuse
    first_rPr = None
    first_run = elem.find(qn("r"))
    if first_run is not None:
        rPr = first_run.find(qn("rPr"))
        if rPr is not None:
            first_rPr = copy.deepcopy(rPr)

    # Remove everything except pPr
    for child in list(elem):
        tag = child.tag if isinstance(child.tag, str) else ""
        if tag == qn("pPr"):
            continue
        elem.remove(child)

    # Create new run with text
    run = etree.SubElement(elem, qn("r"))
    if first_rPr is not None:
        run.insert(0, first_rPr)
    t = etree.SubElement(run, qn("t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def remove_paragraph(para):
    """Remove a paragraph element from the document body."""
    elem = para._element
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)


def insert_paragraph_after(doc, ref_para, text, style_name=None, copy_rPr_from=None):
    """
    Insert a new paragraph after ref_para with given text.
    If style_name is given, set the style.
    If copy_rPr_from is a paragraph, copy first run's rPr.
    Returns the new paragraph element.
    """
    body = doc.element.body
    ref_elem = ref_para._element

    # Create new w:p
    new_p = etree.Element(qn("p"))
    pPr = etree.SubElement(new_p, qn("pPr"))

    if style_name:
        pStyle = etree.SubElement(pPr, qn("pStyle"))
        pStyle.set(qn("val"), style_name)

    # Copy paragraph formatting from reference if same style
    if copy_rPr_from is not None:
        ref_pPr = copy_rPr_from._element.find(qn("pPr"))
        if ref_pPr is not None:
            for child in ref_pPr:
                tag = child.tag if isinstance(child.tag, str) else ""
                if tag != qn("pStyle"):
                    pPr.append(copy.deepcopy(child))

    # Create run
    run = etree.SubElement(new_p, qn("r"))

    # Copy rPr from reference paragraph's first run
    if copy_rPr_from is not None:
        first_run = copy_rPr_from._element.find(qn("r"))
        if first_run is not None:
            rPr = first_run.find(qn("rPr"))
            if rPr is not None:
                run.insert(0, copy.deepcopy(rPr))

    t = etree.SubElement(run, qn("t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Insert after ref
    ref_elem.addnext(new_p)
    return new_p


def insert_empty_paragraph_after(doc, ref_para, style_name=None):
    """Insert an empty paragraph after ref_para."""
    body = doc.element.body
    ref_elem = ref_para._element
    new_p = etree.Element(qn("p"))
    if style_name:
        pPr = etree.SubElement(new_p, qn("pPr"))
        pStyle = etree.SubElement(pPr, qn("pStyle"))
        pStyle.set(qn("val"), style_name)
    ref_elem.addnext(new_p)
    return new_p


def set_table_cell_text(cell, text):
    """Set text of a table cell, preserving first paragraph formatting."""
    # Clear all paragraphs except first
    for para in cell.paragraphs[1:]:
        remove_paragraph(para)
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)


def clear_para(para):
    """Clear a paragraph to empty, removing all runs."""
    elem = para._element
    for child in list(elem):
        tag = child.tag if isinstance(child.tag, str) else ""
        if tag != qn("pPr"):
            elem.remove(child)


def get_para_by_index(doc, idx):
    """Get paragraph by index, with bounds check."""
    if 0 <= idx < len(doc.paragraphs):
        return doc.paragraphs[idx]
    return None


# ============================================================================
# MAIN TRANSFORMATION
# ============================================================================

def transform_template():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))

    paras = doc.paragraphs
    total = len(paras)
    print(f"Template has {total} paragraphs, {len(doc.tables)} tables")

    # ========================================================================
    # PAGE 1 — Cover page (paragraphs 0-26)
    # Structure for juristische Person:
    #   "Gutachten" (not "Gutachten und Bericht")
    #   "in dem Insolvenzantragsverfahren über das Vermögen des/der"
    #   KI_Schuldner_Firma
    #   KI_Schuldner_Betriebsstaette
    #   vertreten durch
    #   KI_SCHULDNER_GESCHAEFTSFUEHRER
    #   (Geschäftsführer)
    #   - nachfolgend "Antragsgegnerin" -
    #   Amtsgericht KI_Gericht_Ort
    #   -KI_Akte_GerichtAZ-
    # ========================================================================

    # Para 5: "Gutachten und Bericht" → just "Gutachten"
    set_paragraph_text(paras[5], "Gutachten")

    # Para 7: fix article to "des/der" for genitive
    set_paragraph_text(paras[7], "in dem Insolvenzantragsverfahren über das Vermögen KI_Schuldner_des_der")

    # Para 9: was KI_Akte_Bezeichnung → replace with Firma
    set_paragraph_text(paras[9], "KI_Schuldner_Firma")
    # Para 10: was KI_Schuldner_Adr → replace with Betriebsstaette
    set_paragraph_text(paras[10], "KI_Schuldner_Betriebsstaette")

    # Para 12: "vertreten durch" — keep as is
    # Para 13: KI_SCHULDNER_GESCHAEFTSFUEHRER — keep as is
    # Para 14 is empty — set to "(Geschäftsführer)"
    set_paragraph_text(paras[14], "(Geschäftsführer)")

    # Para 15: empty → set to "- nachfolgend „Antragsgegnerin" -"
    set_paragraph_text(paras[15], '- nachfolgend \u201eAntragsgegnerin\u201c -')

    # Clear paras 16-24 (spacers before Amtsgericht)
    for idx in range(16, 25):
        clear_para(paras[idx])

    # Para 25: "Amtsgericht KI_Gericht_Ort" — keep as is
    # Para 26: fix AZ formatting
    set_paragraph_text(paras[26], "-KI_Akte_GerichtAZ-")

    # ========================================================================
    # PAGE 2 — Sachverständiger + Ergebnis (paragraphs 28-60)
    # Keep Sachverständiger block (28-38) as is — already correct.
    # Add Datum line. Replace result summary.
    # ========================================================================

    # Para 41: empty → add Datum line
    set_paragraph_text(paras[41], "                                          Datum: \t[[SLOT_001: Datum Gutachten]]")

    # Para 43: "welches zu dem Ergebnis kommt:" — keep

    # Replace result placeholder paragraphs (46-60)
    set_paragraph_text(paras[46], "[[SLOT_002: Ergebnis Punkt 1 - Feststellung Zahlungsunfähigkeit]]")
    set_paragraph_text(paras[48], "[[SLOT_003: Ergebnis Punkt 2 - Feststellung Überschuldung]]")
    set_paragraph_text(paras[50], "[[SLOT_004: Ergebnis Punkt 3 - Verfahrenskostendeckung]]")
    set_paragraph_text(paras[52], "[[SLOT_005: Ergebnis Punkt 4 - Empfehlung Sicherungsmaßnahmen/Eröffnung]]")

    # Clear remaining old result text
    for idx in [54, 55, 56, 57, 58, 59]:
        if idx < total:
            set_paragraph_text(paras[idx], "")
            clear_para(paras[idx])

    # Para 60: "Inhaltsverzeichnis" — keep
    # ========================================================================
    # PAGE 3-4 — Inhaltsverzeichnis (paragraphs 61+)
    # TOC is regenerated by Word. Leave as is.
    # ========================================================================

    # ========================================================================
    # PAGE 5 — A. Gutachtenauftrag / I. Auftrag
    # Para 62: Heading 1 "Gutachtenauftrag und Grundlagen..." — keep
    # Para 65: Heading 2 "Auftrag" — keep
    # ========================================================================

    # Para 67: Auftrag text — update to match JP structure
    set_paragraph_text(paras[67],
        "Das Amtsgericht KI_Gericht_Ort hat KI_Verwalter_den_die Unterzeichner mit Beschluss vom "
        "KI_Akte_BeschlussDat beauftragt, in dem")

    # Para 69: "Insolvenzantragsverfahren..." — fix to use des_der
    set_paragraph_text(paras[69], "Insolvenzantragsverfahren über das Vermögen KI_Schuldner_des_der")

    # Para 71: was KI_Akte_Bezeichnung → Schuldner Firma block
    set_paragraph_text(paras[71], "KI_Schuldner_Firma")
    # Para 72: empty → Betriebsstaette
    set_paragraph_text(paras[72], "KI_Schuldner_Betriebsstaette")
    # Para 73: empty → "vertreten durch"
    set_paragraph_text(paras[73], "")
    # Para 74: empty → GF block
    set_paragraph_text(paras[74], "vertreten durch")

    # Insert the rest of the Auftrag Schuldner block after para 74
    ref = paras[74]._element
    lines_to_insert = [
        "KI_SCHULDNER_GESCHAEFTSFUEHRER",
        "(Geschäftsführer)",
        "",
        '- nachfolgend \u201eAntragsgegnerin\u201c -',
        "",
        "ein schriftliches Gutachten zur Beantwortung folgender Fragen zu erstellen:",
        "",
        "a) Liegen Tatsachen vor, wonach der Schluss auf (drohende) Zahlungsunfähigkeit der Antragsgegnerin gerechtfertigt ist?",
        "",
        "Falls ja.",
        "",
        "b) Ist eine die Verfahrenskosten (§ 54 InsO) deckende Masse vorhanden? Dabei sind auch insolvenzspezifische Ansprüche (Haftung von verantwortlichen Organen, Anfechtungsansprüche) zu prüfen und darzustellen, wann die Zahlungsunfähigkeit / Überschuldung eingetreten ist.",
        "",
        "c) Erscheinen vorläufige Anordnungen zur Sicherung der Masse (allgemeines Veräußerungsverbot, vorläufige Verwaltung, Postsperre usw.) erforderlich?",
        "",
    ]

    ref_pPr = paras[71]._element.find(qn("pPr"))
    for line_text in lines_to_insert:
        new_p = etree.Element(qn("p"))
        if ref_pPr is not None:
            new_p.append(copy.deepcopy(ref_pPr))
        if line_text:
            run = etree.SubElement(new_p, qn("r"))
            t = etree.SubElement(run, qn("t"))
            t.text = line_text
            if line_text[0] == " " or line_text[-1] == " ":
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        ref.addnext(new_p)
        ref = new_p

    # Para 75: was "ein schriftliches Gutachten..." → clear (we inserted it above)
    set_paragraph_text(paras[75], "")
    clear_para(paras[75])
    set_paragraph_text(paras[76], "")
    clear_para(paras[76])
    set_paragraph_text(paras[77], "")
    clear_para(paras[77])
    set_paragraph_text(paras[78], "")
    clear_para(paras[78])

    # Para 79: "In Erfüllung dieses Auftrages..." — update
    set_paragraph_text(paras[79],
        "In Erfüllung dieses Auftrages hat KI_Verwalter_der_die Unterzeichner das nachfolgende "
        "Gutachten erstellt und erstattet hiermit Bericht über den Verlauf des Antragsverfahrens:")

    # ========================================================================
    # PAGE 6 — II. Grundlagen (paragraphs 80+)
    # ========================================================================

    # Para 80: Heading 2 "Grundlagen" — keep

    # Para 82: Grundlagen intro — update for JP
    set_paragraph_text(paras[82],
        "Dem Beschluss ging ein [[SLOT_006: Eigenantrag/Fremdantrag]] KI_Schuldner_des_der "
        "Antragsgegnerin vom KI_Akte_VorlVerfahrDat voraus, welcher am "
        "[[SLOT_007: Eingangsdatum bei Gericht]] bei dem Amtsgericht KI_Gericht_Ort einging.")

    # Insert second paragraph of Grundlagen after 82
    ref = paras[82]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)  # empty line
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = ("Dieses Gutachten basiert auf der Auswertung der Gerichtsakten, Gesprächen mit den unten "
              "genannten Auskunftspersonen, Inaugenscheinnahmen, den Ergebnissen des beauftragten "
              "Sachverständigen und den von der Antragsgegnerin vorgelegten Unterlagen.")
    ref.addnext(new_p)

    # Para 84: Heading 3 "Unterlagen" — keep

    # Replace Unterlagen content (paras 86-95)
    set_paragraph_text(paras[86],
        "Das Insolvenzgericht hat KI_Verwalter_dem_der Unterzeichner am "
        "[[SLOT_008: Datum Aktenübersendung]] die Gerichtsakte übersandt.")

    set_paragraph_text(paras[88],
        "[[SLOT_009: Beschreibung des ersten Besprechungstermins mit dem Geschäftsführer]]")

    set_paragraph_text(paras[90],
        "Dem Unterzeichner standen darüber hinaus folgende Unterlagen zur Einsichtnahme zur Verfügung:")

    set_paragraph_text(paras[92],
        "[[SLOT_010: Liste der Unterlagen (Buchungskonten, OPOS-Listen, Handelsregisterauszug, Gesellschaftsvertrag, BWA, Jahresabschlüsse etc.)]]")

    # Clear old datensicherung text
    set_paragraph_text(paras[94], "")
    clear_para(paras[94])

    # Para 96: Heading 3 "Auskunftspersonen" — keep

    # Para 98: was empty → replace with intro
    set_paragraph_text(paras[98],
        "Darüber hinaus standen KI_Verwalter_dem_der Unterzeichner folgende Personen für Auskünfte zur Verfügung:")

    # Insert Auskunftspersonen list slot after 98
    ref = paras[98]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = "[[SLOT_011: Liste der Auskunftspersonen mit Name, Funktion und Organisation]]"
    ref.addnext(new_p)

    # Para 100: Heading 2 "VID Erklärung" — keep
    # Paras 102-108: VID boilerplate — KEEP AS IS (already correct)

    # ========================================================================
    # PAGE 8-9 — B. Allgemeine Angaben / Tables
    # ========================================================================

    # Para 111: Heading 1 "Allgemeine Angaben" — keep
    # Para 114: Heading 2 "Angaben zur Schuldnerin..." — keep
    # Para 116: Heading 3 "Statistische Angaben" — keep

    # Table 0: Statistische Angaben (14 rows x 2 cols)
    t0 = doc.tables[0]

    # Row 0: Firma → KI_Schuldner_Firma
    set_table_cell_text(t0.rows[0].cells[0], "Firma")
    set_table_cell_text(t0.rows[0].cells[1], "KI_Schuldner_Firma")

    # Row 1: Rechtsform → KI_Schuldner_Rechtsform
    set_table_cell_text(t0.rows[1].cells[0], "Rechtsform")
    set_table_cell_text(t0.rows[1].cells[1], "KI_Schuldner_Rechtsform")

    # Row 2: Satzungsgemäßer Sitz
    set_table_cell_text(t0.rows[2].cells[0], "Satzungsgemäßer Sitz")
    set_table_cell_text(t0.rows[2].cells[1], "[[SLOT_012: Satzungssitz laut Gesellschaftsvertrag]]")

    # Row 3: Verwaltungssitz
    set_table_cell_text(t0.rows[3].cells[0], "Verwaltungssitz")
    set_table_cell_text(t0.rows[3].cells[1], "KI_Schuldner_Betriebsstaette")

    # Row 4: Handelsregister
    set_table_cell_text(t0.rows[4].cells[0], "Handelsregister")
    set_table_cell_text(t0.rows[4].cells[1], "KI_Schuldner_HRB")

    # Row 5: Insolvenzforderungen
    set_table_cell_text(t0.rows[5].cells[0], "Insolvenzforderungen, davon\ngesichert\nnachrangig")
    set_table_cell_text(t0.rows[5].cells[1],
        "KI_Forderungen_Gesamt\nKI_Forderungen_Gesichert\n[[SLOT_013: Nachrangige Forderungen]]")

    # Row 6: Antragstellerin
    set_table_cell_text(t0.rows[6].cells[0], "Antragstellerin")
    set_table_cell_text(t0.rows[6].cells[1], "KI_Antragsteller_Name")

    # Row 7: Antragsgrund
    set_table_cell_text(t0.rows[7].cells[0], "Antragsgrund")
    set_table_cell_text(t0.rows[7].cells[1], "KI_Verfahren_Eroeffnungsgrund")

    # Row 8: Internationaler Bezug
    set_table_cell_text(t0.rows[8].cells[0], "Internationaler Bezug")
    set_table_cell_text(t0.rows[8].cells[1], "[[SLOT_014: Internationaler Bezug ja/nein mit Erläuterung]]")

    # Row 9: Eigenverwaltung
    set_table_cell_text(t0.rows[9].cells[0], "Eigenverwaltung")
    set_table_cell_text(t0.rows[9].cells[1], "[[SLOT_015: Eigenverwaltung ja/nein]]")

    # Row 10: Geschäftszweig
    set_table_cell_text(t0.rows[10].cells[0], "Geschäftszweig")
    set_table_cell_text(t0.rows[10].cells[1], "[[SLOT_016: Geschäftszweig/Branche]]")

    # Row 11: Unternehmensgegenstand
    set_table_cell_text(t0.rows[11].cells[0], "Unternehmensgegenstand\ngemäß Handelsregister")
    set_table_cell_text(t0.rows[11].cells[1], "[[SLOT_017: Unternehmensgegenstand laut HR-Eintragung]]")

    # Row 12: Arbeitnehmer
    set_table_cell_text(t0.rows[12].cells[0], "Arbeitnehmer")
    set_table_cell_text(t0.rows[12].cells[1], "KI_Arbeitnehmer_Anzahl")

    # Row 13: Betriebsrat
    set_table_cell_text(t0.rows[13].cells[0], "Betriebsrat")
    set_table_cell_text(t0.rows[13].cells[1], "[[SLOT_018: Betriebsrat ja/nein]]")

    # Para 119: Heading 3 "Gesellschaftsrechtliche Angaben" — keep

    # Table 1: Gesellschaftsrechtliche Angaben (9 rows x 2 cols)
    t1 = doc.tables[1]

    # Row 0: Stammkapital
    set_table_cell_text(t1.rows[0].cells[0], "Stammkapital (GmbH) /\nGrundkapital (AG)")
    set_table_cell_text(t1.rows[0].cells[1], "[[SLOT_019: Stammkapital Betrag in EUR]]")

    # Row 1: Gesellschafter
    set_table_cell_text(t1.rows[1].cells[0], "Gesellschafter (GmbH) /\nAktionäre (AG)")
    set_table_cell_text(t1.rows[1].cells[1], "KI_Gesellschafter")

    # Row 2: Geschäftsführer
    set_table_cell_text(t1.rows[2].cells[0], "Geschäftsführer (GmbH) /\nVorstand (AG)")
    set_table_cell_text(t1.rows[2].cells[1], "[[SLOT_020: Geschäftsführer Name, geb. Datum, Bestellung seit, Details zu § 181 BGB-Befreiung]]")

    # Row 3: Befreiung § 181 BGB
    set_table_cell_text(t1.rows[3].cells[0], "Befreiung von § 181 BGB")
    set_table_cell_text(t1.rows[3].cells[1], "[[SLOT_021: Befreiung von § 181 BGB ja/nein]]")

    # Row 4: D&O Versicherung
    set_table_cell_text(t1.rows[4].cells[0], "D&O Versicherung")
    set_table_cell_text(t1.rows[4].cells[1], "[[SLOT_022: D&O Versicherung ja/nein mit Versicherungssumme]]")

    # Row 5: Größenklasse
    set_table_cell_text(t1.rows[5].cells[0], "Größenklasse gem. HGB")
    set_table_cell_text(t1.rows[5].cells[1], "KI_Groessenklasse")

    # Row 6: Gründung
    set_table_cell_text(t1.rows[6].cells[0], "Gründung")
    set_table_cell_text(t1.rows[6].cells[1], "KI_Gruendung")

    # Row 7: Eintragung HR
    set_table_cell_text(t1.rows[7].cells[0], "Eintragung ins Handelsregister")
    set_table_cell_text(t1.rows[7].cells[1], "KI_HR_Eintragung")

    # Row 8: wirtschaftlich Berechtigter
    set_table_cell_text(t1.rows[8].cells[0], "wirtschaftlich Berechtigter\nnach Transparenzregister")
    set_table_cell_text(t1.rows[8].cells[1], "[[SLOT_023: Wirtschaftlich Berechtigter nach Transparenzregister]]")

    # Para 122: Heading 3 "Steuerrechtliche Angaben" — keep

    # Table 2: Steuerrechtliche Angaben (5 rows x 2 cols)
    t2 = doc.tables[2]

    # Row 0: Finanzamt
    set_table_cell_text(t2.rows[0].cells[0], "Finanzamt")
    set_table_cell_text(t2.rows[0].cells[1], "KI_Finanzamt")

    # Row 1: Steuer-Nr.
    set_table_cell_text(t2.rows[1].cells[0], "Steuer-Nr. / USt-ID-Nr.")
    set_table_cell_text(t2.rows[1].cells[1], "KI_Steuernummer / [[SLOT_024: USt-ID-Nr.]]")

    # Row 2: Wirtschaftsjahr / USt
    set_table_cell_text(t2.rows[2].cells[0], "Wirtschaftsjahr\nUSt-Versteuerung")
    set_table_cell_text(t2.rows[2].cells[1], "[[SLOT_025: Wirtschaftsjahr]] / [[SLOT_026: Soll- oder Istversteuerung]]")

    # Row 3: Steuerliche Organschaft
    set_table_cell_text(t2.rows[3].cells[0], "Steuerliche Organschaft")
    set_table_cell_text(t2.rows[3].cells[1], "[[SLOT_027: Steuerliche Organschaft ja/nein]]")

    # Row 4: Letzter Jahresabschluss
    set_table_cell_text(t2.rows[4].cells[0], "Letzter Jahresabschluss")
    set_table_cell_text(t2.rows[4].cells[1], "KI_Letzter_Jahresabschluss")

    # Para 125: Heading 3 "Sonstige Angaben" — keep

    # Table 3: Sonstige Angaben (10 rows x 2 cols)
    t3 = doc.tables[3]

    # Row 0: Telefon & E-Mail
    set_table_cell_text(t3.rows[0].cells[0], "Telefon & E-Mail\nGeschäftsführer")
    set_table_cell_text(t3.rows[0].cells[1], "[[SLOT_028: Telefon und E-Mail des Geschäftsführers]]")

    # Row 1: Sozialversicherungsträger
    set_table_cell_text(t3.rows[1].cells[0], "Sozialversicherungsträger")
    set_table_cell_text(t3.rows[1].cells[1], "KI_SVTraeger")

    # Row 2: Betriebsnummer
    set_table_cell_text(t3.rows[2].cells[0], "Betriebsnummer")
    set_table_cell_text(t3.rows[2].cells[1], "KI_Betriebsnummer")

    # Row 3: Berufsgenossenschaft
    set_table_cell_text(t3.rows[3].cells[0], "Berufsgenossenschaft\nMitgliedsnummer")
    set_table_cell_text(t3.rows[3].cells[1], "[[SLOT_029: Berufsgenossenschaft und Mitgliedsnummer]]")

    # Row 4: älteste Forderung
    set_table_cell_text(t3.rows[4].cells[0], "älteste Forderung")
    set_table_cell_text(t3.rows[4].cells[1], "KI_Aelteste_Forderung")

    # Row 5: Zeitpunkt erste bilanzielle Überschuldung
    set_table_cell_text(t3.rows[5].cells[0], "Zeitpunkt erste bilanzielle\nÜberschuldung")
    set_table_cell_text(t3.rows[5].cells[1], "[[SLOT_030: Zeitpunkt erste bilanzielle Überschuldung]]")

    # Row 6: Steuerberater
    set_table_cell_text(t3.rows[6].cells[0], "Steuerberater")
    set_table_cell_text(t3.rows[6].cells[1], "KI_Steuerberater")

    # Row 7: Bankverbindungen
    set_table_cell_text(t3.rows[7].cells[0], "Bankverbindungen")
    set_table_cell_text(t3.rows[7].cells[1], "KI_Bankverbindungen")

    # Row 8: Insolvenzsonderkonto
    set_table_cell_text(t3.rows[8].cells[0], "Insolvenzsonderkonto")
    set_table_cell_text(t3.rows[8].cells[1], "KI_ANDERKONTO_IBAN / KI_Anderkonto_Bank")

    # Row 9: Gerichtsvollzieher
    set_table_cell_text(t3.rows[9].cells[0], "Zuständiger\nGerichtsvollzieher")
    set_table_cell_text(t3.rows[9].cells[1], "KI_Gerichtsvollzieher")

    # ========================================================================
    # B.II. Angaben zum Verfahren (para 128+)
    # Structure: 1. Verfahrenshistorie, 2. Örtliche Zuständigkeit
    # ========================================================================

    # Para 128: Heading 2 "Angaben zum Verfahren" — keep

    # Para 130: empty → Verfahrenshistorie slot
    set_paragraph_text(paras[130],
        "[[SLOT_031: Verfahrenshistorie - Antragsdatum, Eingangsdatum, Beschlussdatum, "
        "Bestellung, Sicherungsmaßnahmen, Gutachtenauftrag]]")

    # Para 132: Heading 3 "Internationale Zuständigkeit..." — keep
    # Paras 134-138: International jurisdiction boilerplate — keep as is
    # Para 140: empty → keep

    # Para 142: Heading 3 "Örtliche Zuständigkeit..." — keep
    # Para 144: Örtliche Zuständigkeit boilerplate — keep

    # Para 146: empty → concrete jurisdiction reasoning
    set_paragraph_text(paras[146],
        "[[SLOT_032: Konkrete Zuständigkeitsbegründung - wo ist der Verwaltungssitz/Geschäftsbetrieb, Gerichtsbezirk]]")

    # Para 148: empty → conclusion
    set_paragraph_text(paras[148],
        "Das Amtsgericht KI_Gericht_Ort ist daher örtlich zuständig.")

    # Para 150: Heading 3 "Sicherungsmaßnahmen" — keep
    # Para 152: empty → Sicherungsmaßnahmen content
    set_paragraph_text(paras[152],
        "[[SLOT_033: Sicherungsmaßnahmen - Beschluss, Datum, Maßnahmen "
        "(Zustimmungsvorbehalt, Zwangsvollstreckungsverbot, Kontensperre, vorl. Gläubigerausschuss)]]\n")

    # Para 154: was SLOT_105 → clear (merged into 152)
    set_paragraph_text(paras[154], "")
    clear_para(paras[154])

    # Para 156: Heading 3 "Einzelermächtigungen" — keep
    # Para 158: empty → Einzelermächtigungen content
    set_paragraph_text(paras[158],
        "[[SLOT_034: Einzelermächtigungen des vorläufigen Insolvenzverwalters]]")

    # ========================================================================
    # C. Informationen über das Unternehmen (para 161+)
    # ========================================================================

    # Para 161: Heading 1 "Informationen über das Unternehmen" — keep

    # C.I. Unternehmensgegenstand (para 164)
    # Para 164: Heading 2 — keep
    # Para 166: empty → narrative slot
    set_paragraph_text(paras[166],
        "[[SLOT_035: Unternehmensgegenstand - Geschäftstätigkeit, Branche, "
        "wesentliche Produkte/Dienstleistungen]]")

    # C.II. Gesellschaftsrechtliche Verhältnisse (para 168)
    # Para 168: Heading 2 — keep
    # Para 170: empty → narrative
    set_paragraph_text(paras[170],
        "[[SLOT_036: Gesellschaftsrechtliche Verhältnisse - Gründung, Gesellschafter, "
        "Stammkapital, Geschäftsführer, Veränderungen, Gesellschafterstruktur-Tabelle]]")

    # C.III. Finanzierungsstruktur (para 172)
    # Para 172: Heading 2 — keep
    # Para 174: empty → narrative
    set_paragraph_text(paras[174],
        "[[SLOT_037: Finanzierungsstruktur - Bankverbindungen mit Kreditlinien, "
        "Darlehen, Kontokorrent, Sicherheiten pro Bank, Leasing, sonstige Finanzierungen]]")

    # Clear old paras 175-177
    for idx in [175, 176, 177]:
        set_paragraph_text(paras[idx], "")
        clear_para(paras[idx])

    # C.IV. Wesentliche Vertragsverhältnisse (para 178)
    # Para 178: Heading 2 "Wesentliche Vertragsverhältnisse" — keep

    # 1. Arbeitsrechtliche Verhältnisse (para 180)
    # Para 180: Heading 3 — keep
    set_paragraph_text(paras[182],
        "[[SLOT_038: Arbeitsrechtliche Verhältnisse - Anzahl Arbeitnehmer, "
        "Auszubildende, Betriebsrat, Tarifbindung, Kündigungen, Sozialplan]]")
    set_paragraph_text(paras[184],
        "[[SLOT_039: Geschäftsführeranstellungsvertrag - Konditionen, Kündigungsstatus]]")
    # Clear old SLOT_113
    set_paragraph_text(paras[186], "")
    clear_para(paras[186])

    # 2. Gewerberaummietverträge (para 188)
    # Para 188: Heading 3 — keep
    set_paragraph_text(paras[190],
        "[[SLOT_040: Gewerberaummietverträge / Pachtverträge - Objekt, Vermieter, "
        "Miete, Laufzeit, Kündigungsstatus]]")

    # 3. Versicherungsverträge (para 192)
    # Para 192: Heading 3 — keep
    set_paragraph_text(paras[194],
        "[[SLOT_041: Versicherungsverträge - Art, Versicherer, Beiträge, Status]]")

    # 4. Versorgungsverträge (para 196)
    # Para 196: Heading 3 — keep
    set_paragraph_text(paras[198],
        "[[SLOT_042: Versorgungsverträge - Strom, Gas, Wasser, Telekommunikation]]")

    # 5. Leasingverträge (para 200)
    # Para 200: Heading 3 — keep
    set_paragraph_text(paras[202],
        "[[SLOT_043: Leasingverträge / Mietverträge über bewegliche Gegenstände]]")

    # 6. Factoring (para 204)
    # Para 204: Heading 3 — keep
    set_paragraph_text(paras[206],
        "[[SLOT_044: Factoring-Vereinbarungen oder 'Factoring-Vereinbarungen bestehen nicht.']]")

    # 7. Sonstige Verträge (para 208)
    # Para 208: Heading 3 — keep
    set_paragraph_text(paras[210],
        "[[SLOT_045: Sonstige Verträge oder 'Sonstige wesentliche Vertragsverhältnisse sind nicht bekannt.']]")

    # C.V. Arbeitsrechtliche Verhältnisse (broader section, para 212)
    # Para 212: Heading 2 "Arbeitsrechtliche Verhältnisse" — keep (for larger companies)
    set_paragraph_text(paras[214],
        "[[SLOT_046: Erweiterte arbeitsrechtliche Verhältnisse - "
        "kollektivarbeitsrechtliche und betriebsverfassungsrechtliche Aspekte]]")
    # Clear old slots
    set_paragraph_text(paras[216], "")
    clear_para(paras[216])
    set_paragraph_text(paras[218], "")
    clear_para(paras[218])

    # C.VI. Zahlungsverkehr (para 220)
    # Para 220: Heading 2 — keep
    set_paragraph_text(paras[222],
        "Im Zeitpunkt der Antragstellung verfügte die Antragsgegnerin über nachstehende Girokonten:")

    # Table 4 (Konten) — update with generic slots
    t4 = doc.tables[4]
    # Header row
    set_table_cell_text(t4.rows[0].cells[0], "Konto")
    set_table_cell_text(t4.rows[0].cells[1], "Stand bei Antragstellung")
    # Data rows — update existing with slots
    set_table_cell_text(t4.rows[1].cells[0], "")
    set_table_cell_text(t4.rows[1].cells[1], "")
    set_table_cell_text(t4.rows[2].cells[0], "[[SLOT_047: Kreditinstitut 1 mit Kontonummer]]")
    set_table_cell_text(t4.rows[2].cells[1], "[[SLOT_048: Saldo 1]] EUR")
    set_table_cell_text(t4.rows[3].cells[0], "")
    set_table_cell_text(t4.rows[3].cells[1], "")
    set_table_cell_text(t4.rows[4].cells[0], "[[SLOT_049: Kreditinstitut 2 mit Kontonummer]]")
    set_table_cell_text(t4.rows[4].cells[1], "[[SLOT_050: Saldo 2]] EUR")
    # Clear remaining rows
    for ri in range(5, len(t4.rows)):
        set_table_cell_text(t4.rows[ri].cells[0], "")
        set_table_cell_text(t4.rows[ri].cells[1], "")

    # Para 227: Kontensperrung text — update for JP
    set_paragraph_text(paras[227],
        "Die Konten wurden für Verfügungen KI_Schuldner_des_der Antragsgegnerin gesperrt. "
        "Seit Anordnung der vorläufigen Verwaltung werden Zahlungseingänge über das "
        "Insolvenz-Sonderkonto geleitet.")

    # C.VII. Rechnungswesen (para 229)
    # Para 229: Heading 2 "Rechnungswesen" — keep

    # Para 231: Heading 3 "Buchhaltung vor Antragstellung" — keep
    set_paragraph_text(paras[233],
        "[[SLOT_051: Buchhaltung vor Antragstellung - wer hat geführt, Qualität, Vollständigkeit, "
        "letzter Jahresabschluss]]")
    # Clear old text
    set_paragraph_text(paras[235], "")
    clear_para(paras[235])

    # Para 237: Heading 3 "Buchhaltung nach Antragstellung" — keep
    set_paragraph_text(paras[239],
        "[[SLOT_052: Buchhaltung nach Antragstellung - Fortführung durch wen, Maßnahmen]]")

    # C.VIII. Anhängige Rechtsstreitigkeiten (para 241)
    # Para 241: Heading 2 — keep
    set_paragraph_text(paras[243],
        "[[SLOT_053: Anhängige Rechtsstreitigkeiten oder "
        "'Anhängige Rechtsstreitigkeiten sind dem Unterzeichner nicht bekannt.']]")

    # C.IX. Betriebswirtschaftliche Verhältnisse (para 245)
    # Para 245: Heading 2 — keep

    # 1. Wirtschaftliche Entwicklung und Krisenursache (para 247)
    # Para 247: Heading 3 — keep
    set_paragraph_text(paras[249],
        "[[SLOT_054: Wirtschaftliche Entwicklung und Krisenursache - "
        "ausführliche Darstellung der Unternehmensentwicklung, Umsatzentwicklung, "
        "Krisenursachen, ggf. mit Jahresabschlusszahlen]]")

    # 2. Maßnahmen im Insolvenzeröffnungsverfahren (para 251)
    # Para 251: Heading 3 — keep

    # a) Vermögenssicherung (para 253)
    set_paragraph_text(paras[255],
        "[[SLOT_055: Vermögenssicherung - Maßnahmen zur Sicherung des Schuldnervermögens]]")

    # b) Sicherung der Betriebsfortführung (para 257)
    set_paragraph_text(paras[259],
        "[[SLOT_056: Sicherung der Betriebsfortführung - Betriebsfortführung ja/nein, "
        "Maßnahmen, Personal, Aufträge, Lieferanten]]")

    # Clear the old detailed paras 261-278
    for idx in range(261, 279):
        if idx < total:
            set_paragraph_text(paras[idx], "")
            clear_para(paras[idx])

    # Datensicherung (para 269 was "Datensicherung" list para)
    set_paragraph_text(paras[269], "Datensicherung")
    set_paragraph_text(paras[271],
        "[[SLOT_057: Datensicherung - Umfang, Medium, Aufbewahrung]]")

    # 3. Sanierungsaussichten (para 280)
    # Para 280: Heading 3 — keep
    set_paragraph_text(paras[282],
        "[[SLOT_058: Sanierungsaussichten - Analyse der Sanierungsfähigkeit, "
        "ggf. Fortführungsprognose, ggf. Verwertungskonzept]]")

    # ========================================================================
    # D. Vermögensverhältnisse (para 285+)
    # ========================================================================

    # Para 285: Heading 1 "Vermögensverhältnisse" — keep

    # D.I. Bewertungsmaßstab (para 288)
    # Para 288: Heading 2 — keep
    # Para 290: boilerplate — keep as is

    # D.II. Sachverständige Bewertung (para 292)
    # Para 292: Heading 2 — keep
    set_paragraph_text(paras[294],
        "KI_Verwalter_Der_Die_Groß Unterzeichner hat die Bewertung und Inventarisierung "
        "des beweglichen Sachanlagevermögens in Auftrag gegeben.")
    set_paragraph_text(paras[296],
        "Die Aufnahme erfolgte zum Bewertungsstichtag des [[SLOT_059: Bewertungsstichtag]] "
        "durch den Sachverständigen [[SLOT_060: Sachverständiger Name und Adresse]].")

    # Insert Anlage reference after 296
    ref = paras[296]._element
    new_p = etree.Element(qn("p"))
    ref.addnext(new_p)
    ref = new_p
    new_p = etree.Element(qn("p"))
    run = etree.SubElement(new_p, qn("r"))
    t = etree.SubElement(run, qn("t"))
    t.text = ("Inhaltlich nimmt KI_Verwalter_der_die Unterzeichner insoweit Bezug auf das "
              "als Anlage II beigefügte Gutachten vom [[SLOT_061: Datum Wertgutachten]].")
    ref.addnext(new_p)

    # D.III. Aktiva (para 298)
    # Para 298: Heading 2 "Aktiva" — keep

    # 1. Anlagevermögen (para 300)
    # Para 300: Heading 3 — keep

    # a) Immaterielles Vermögen (para 302)
    # aa) Goodwill (304)
    set_paragraph_text(paras[306],
        "[[SLOT_062: Goodwill / Firmenwert Beschreibung und Bewertung]]")

    # bb) selbst geschaffene Schutzrechte (308)
    set_paragraph_text(paras[310],
        "[[SLOT_063: Selbst geschaffene gewerbliche Schutzrechte Beschreibung/Wert]]")

    # cc) entgeltlich erworbene (312)
    set_paragraph_text(paras[314],
        "[[SLOT_064: Entgeltlich erworbene Konzessionen und Schutzrechte Beschreibung/Wert]]")

    # b) Sachanlagen (para 316)
    # aa) Grundstücke (318)
    set_paragraph_text(paras[320],
        "[[SLOT_065: Grundstücke und grundstücksgleiche Rechte - Beschreibung, "
        "Grundbuch, Verkehrswert, Belastungen]]")

    # bb) technische Anlagen (322)
    set_paragraph_text(paras[324],
        "[[SLOT_066: Technische Anlagen und Maschinen - Beschreibung, Bewertung]]")

    # cc) BGA (326)
    set_paragraph_text(paras[328],
        "[[SLOT_067: Betriebs- und Geschäftsausstattung - Beschreibung, Bewertung]]")

    # c) Finanzanlagen (para 330)
    # aa) Anteile verbundene Unternehmen (332)
    set_paragraph_text(paras[334],
        "[[SLOT_068: Anteile an verbundenen Unternehmen Beschreibung/Wert]]")

    # bb) Beteiligungen (336)
    set_paragraph_text(paras[338],
        "[[SLOT_069: Beteiligungen Beschreibung/Wert]]")

    # cc) Wertpapiere Anlagevermögen (340)
    set_paragraph_text(paras[342],
        "[[SLOT_070: Wertpapiere des Anlagevermögens Beschreibung/Wert]]")

    # 2. Umlaufvermögen (para 344)
    # Para 344: Heading 3 — keep

    # a) Vorräte (para 346)
    # aa) Roh-/Hilfs-/Betriebsstoffe (348)
    set_paragraph_text(paras[350],
        "[[SLOT_071: Roh-, Hilfs- und Betriebsstoffe Beschreibung/Wert]]")

    # bb) unfertige Erzeugnisse (352)
    set_paragraph_text(paras[354],
        "[[SLOT_072: Unfertige Erzeugnisse, unfertige Leistungen Beschreibung/Wert]]")

    # cc) fertige Erzeugnisse (356)
    set_paragraph_text(paras[358],
        "[[SLOT_073: Fertige Erzeugnisse / Waren Beschreibung/Wert]]")

    # b) Forderungen und sonstige Vermögensgegenstände (para 360)
    # aa) and bb) — paras 362-369
    set_paragraph_text(paras[362],
        "[[SLOT_074: Forderungen aus Lieferungen und Leistungen - Bestand, Einbringlichkeit, Bewertung]]")
    set_paragraph_text(paras[364], "")
    set_paragraph_text(paras[366],
        "[[SLOT_075: Forderungen gegen verbundene Unternehmen]]")
    set_paragraph_text(paras[368], "")

    # cc) Forderungen Beteiligungsunternehmen (370)
    set_paragraph_text(paras[372],
        "[[SLOT_076: Forderungen gegen Beteiligungsunternehmen]]")

    # dd) sonstige Vermögensgegenstände (374)
    set_paragraph_text(paras[376],
        "[[SLOT_077: Sonstige Vermögensgegenstände Beschreibung/Wert]]")

    # c) Wertpapiere (para 378)
    set_paragraph_text(paras[380],
        "[[SLOT_078: Wertpapiere des Umlaufvermögens Beschreibung/Wert]]")

    # d) Liquide Mittel (para 381)
    set_paragraph_text(paras[383],
        "[[SLOT_079: Kassenbestand]]")
    # Clear old empty paras
    for idx in [384, 385, 386, 387, 388]:
        set_paragraph_text(paras[idx], "")
        clear_para(paras[idx])

    set_paragraph_text(paras[385],
        "[[SLOT_080: Bankguthaben bei Antragstellung mit Details pro Konto]]")

    # Anderkonto text (389) — update
    set_paragraph_text(paras[389],
        "Der Unterzeichner hat bei der KI_Anderkonto_Bank ein Insolvenz-Sonderkonto "
        "mit der IBAN: KI_ANDERKONTO_IBAN eingerichtet. "
        "Das aktuelle Guthaben beträgt [[SLOT_081: Guthaben Anderkonto]] EUR.")

    # Additional accounts (391) — update
    set_paragraph_text(paras[391],
        "[[SLOT_082: Weitere Konten (Festgeld, Rückstellungskonto) mit IBAN und Guthaben]]")

    # 3. Ansprüche auf Kapitalaufbringung und -erhaltung (para 393)
    # Para 393: Heading 3 — keep (JP-specific section)

    # Heading 4 "Ansprüche auf Einzahlung der Stammeinlagen" (para 395) — keep
    # Para 397: boilerplate about § 19 GmbHG — keep as is
    # Para 399-400: Sonderfall verdeckte Sacheinlage — keep as boilerplate
    # Para 402: Heading 4 empty — set to proper heading
    set_paragraph_text(paras[402], "Kapitalerhaltung §§ 30, 31 GmbHG")
    # Paras 404-408: § 30/31 GmbHG boilerplate — keep as is
    # Para 410: Heading 4 "Existenzvernichtender Eingriff" — keep
    # Paras 412-414: § 826 BGB boilerplate — keep as is

    # Add slot for concrete assessment after boilerplate
    set_paragraph_text(paras[399],
        "[[SLOT_083: Konkrete Prüfung Einzahlung Stammeinlagen - vollständig eingezahlt ja/nein, ggf. Nachschusspflichten]]")
    set_paragraph_text(paras[400], "")
    clear_para(paras[400])
    set_paragraph_text(paras[401], "")
    clear_para(paras[401])

    # Para 406: empty → slot for concrete assessment §§ 30, 31
    set_paragraph_text(paras[406],
        "[[SLOT_084: Konkrete Prüfung §§ 30, 31 GmbHG - verbotene Auszahlungen, Erkenntnisse]]")

    # 4. Zusammenfassung Aktiva (para 416)
    # Para 416: Heading 3 — keep

    # a) Liquidationswerte (418)
    set_paragraph_text(paras[419],
        "[[SLOT_085: Zusammenfassung Liquidationswerte mit Tabelle]]")
    set_paragraph_text(paras[420], "")
    clear_para(paras[420])

    # b) Fortführungswerte (422)
    set_paragraph_text(paras[423],
        "[[SLOT_086: Zusammenfassung Fortführungswerte]]")

    # c) Aussonderung (426)
    set_paragraph_text(paras[428],
        "[[SLOT_087: Aussonderungsrechte Beschreibung]]")
    set_paragraph_text(paras[429],
        "[[SLOT_088: Tabelle Aussonderung]]")

    # d) Absonderung (431)
    set_paragraph_text(paras[433],
        "[[SLOT_089: Absonderungsrechte Beschreibung]]")
    set_paragraph_text(paras[434],
        "[[SLOT_090: Tabelle Absonderung]]")

    # Tables 5 & 6 (Aktiva summary tables) — leave structure, filled dynamically

    # D.IV. Passiva (para 436)
    # Para 436: Heading 2 — keep
    set_paragraph_text(paras[438],
        "[[SLOT_091: Passiva Übersicht - Forderungsanmeldungen nach Kategorien]]")
    set_paragraph_text(paras[440],
        "[[SLOT_092: Passiva-Tabelle mit Forderungsübersicht]]")
    set_paragraph_text(paras[444],
        "[[SLOT_093: Weitere Passiva-Details oder Nachrangige Forderungen]]")

    # D.V. Drittsicherheiten (para 446)
    # Para 446: Heading 2 — keep
    set_paragraph_text(paras[448],
        "[[SLOT_094: Drittsicherheiten - Bürgschaften, Patronatserklärungen, "
        "sonstige Sicherheiten Dritter]]")

    # ========================================================================
    # E. Eröffnungsgründe (para 451+)
    # ========================================================================

    # Para 451: Heading 1 "Eröffnungsgründe" — keep

    # E.I. Zahlungsunfähigkeit (para 454)
    # Para 454: Heading 2 — keep
    # Para 456: Heading 3 "Definition" — keep
    # Para 458: BGH definition boilerplate — KEEP AS IS

    # Para 460: Heading 3 "Vorliegen im Gutachtenfall" — keep
    set_paragraph_text(paras[462],
        "[[SLOT_095: Vorliegen der Zahlungsunfähigkeit im konkreten Fall - "
        "Finanzstatus, Liquiditätslücke, fällige Verbindlichkeiten vs. verfügbare Mittel]]")

    # E.II. Überschuldung (para 464) — relevant for juristische Person!
    # Para 464: Heading 2 — keep
    # Para 466: Heading 3 "Definition" — keep
    # Para 468: § 19 InsO boilerplate — KEEP AS IS (JP-specific, already correct)

    # Para 470: Heading 3 "Vorliegen im Gutachtenfall" — keep
    set_paragraph_text(paras[472],
        "[[SLOT_096: Vorliegen der Überschuldung im konkreten Fall - "
        "Überschuldungsbilanz, Fortführungsprognose, Bewertung]]")

    # E.III. Eintritt der Insolvenzreife (para 474)
    # Para 474: Heading 2 — keep
    set_paragraph_text(paras[476],
        "[[SLOT_097: Eintritt der Insolvenzreife - Zeitpunkt und Begründung]]")
    # Clear old empty placeholders 477-479
    for idx in [477, 478, 479]:
        set_paragraph_text(paras[idx], "")
        clear_para(paras[idx])
    # Para 480: "abschließende Einschätzung" — keep as is

    # ========================================================================
    # F. Insolvenzspezifische Ansprüche (para 483+)
    # ========================================================================

    # Para 483: Heading 1 "Insolvenzspezifische Ansprüche" — keep

    # F.I. Ansprüche aus Insolvenzanfechtung (para 486)
    # Para 486: Heading 2 — keep
    set_paragraph_text(paras[488],
        "[[SLOT_098: Anfechtungsansprüche §§ 129 ff. InsO - "
        "identifizierte anfechtbare Rechtshandlungen, Lookback-Zeiträume, "
        "Erfolgsaussichten und geschätztes Volumen]]")

    # F.II. Ansprüche gegen Geschäftsführer / Vorstand (para 490)
    # Para 490: Heading 2 — keep
    # Para 492: § 15b InsO boilerplate — keep as is
    # Para 496: "Höhe etwaigen Anspruchs..." — keep as is

    # Add slot for concrete assessment
    set_paragraph_text(paras[494],
        "[[SLOT_099: Konkrete Prüfung § 15b InsO - Zahlungen nach Insolvenzreife, "
        "Geschäftsführerhaftung, geschätzte Anspruchshöhe]]")

    # F.III. Kostenbeiträge gem. § 171 InsO (para 499)
    # Para 499: Heading 2 — keep
    set_paragraph_text(paras[501],
        "[[SLOT_100: Kostenbeiträge gem. § 171 InsO - "
        "Sicherungsrechte und daraus resultierende Kostenbeiträge]]")

    # F.IV. Zusammenfassung (para 504)
    # Para 504: Heading 2 — keep
    set_paragraph_text(paras[506],
        "[[SLOT_101: Zusammenfassung insolvenzspezifische Ansprüche - "
        "Gesamtüberblick Anfechtung + GF-Haftung + Kostenbeiträge]]")

    # ========================================================================
    # G. Verfahrenskostendeckung (para 509+)
    # ========================================================================

    # Para 509: Heading 1 — keep
    set_paragraph_text(paras[511],
        "[[SLOT_102: Verfahrenskostendeckung - Berechnung nach InsVV, "
        "Vergütung vorläufiges Verfahren, Vergütung eröffnetes Verfahren, "
        "Gerichtskosten, Gegenüberstellung mit verfügbarer Masse]]")

    # Table 7 (Verfahrenskosten) — update
    t7 = doc.tables[7]
    set_table_cell_text(t7.rows[0].cells[1], "[[SLOT_103: Vergütung vorläufiges Verfahren]] EUR")
    set_table_cell_text(t7.rows[1].cells[1], "[[SLOT_104: Vergütung eröffnetes Verfahren]] EUR")
    set_table_cell_text(t7.rows[2].cells[1], "[[SLOT_105: Gerichtskosten]] EUR")
    set_table_cell_text(t7.rows[3].cells[1], "KI_Verfahrenskosten_Gesamt EUR")

    # Clear remaining paras in G section
    for idx in [512, 513, 514, 515, 516, 517]:
        if idx < total:
            set_paragraph_text(paras[idx], "")
            clear_para(paras[idx])

    # ========================================================================
    # H. Ergebnis und Empfehlung (para 518+)
    # ========================================================================

    # Para 518: Heading 1 — keep
    set_paragraph_text(paras[521],
        "[[SLOT_106: Ergebnis und Empfehlung - Nummerierte Feststellungen "
        "(1. Zahlungsunfähigkeit, 2. Überschuldung, 3. Kostendeckung, 4. Sicherungsmaßnahmen) "
        "und abschließende Empfehlung zur Eröffnung]]")

    # Signature block (para 524+) — keep "Mit freundlichen Grüßen"
    set_paragraph_text(paras[526], "KI_Verwalter_Name")
    set_paragraph_text(paras[527], "als Sachverständiger und vorläufiger Insolvenzverwalter")
    set_paragraph_text(paras[528], "")  # Clear female variant (handled by comment rules)

    # Anlagenverzeichnis (para 531) — keep
    # Table 8 (Anlagen) — update
    t8 = doc.tables[8]
    set_table_cell_text(t8.rows[0].cells[1], "[[SLOT_107: Anlage I Bezeichnung]]")
    set_table_cell_text(t8.rows[1].cells[1], "[[SLOT_108: Anlage II Bezeichnung]]")
    set_table_cell_text(t8.rows[2].cells[1], "[[SLOT_109: Anlage III Bezeichnung]]")
    set_table_cell_text(t8.rows[3].cells[1], "[[SLOT_110: Anlage IV Bezeichnung]]")
    set_table_cell_text(t8.rows[4].cells[1], "[[SLOT_111: Anlage V Bezeichnung]]")

    # ========================================================================
    # SAVE
    # ========================================================================

    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(str(OUTPUT_PATH))
    print("Done!")

    # Verification: reload and check
    print("\n=== VERIFICATION ===")
    doc2 = Document(str(OUTPUT_PATH))
    print(f"Paragraphs: {len(doc2.paragraphs)}")
    print(f"Tables: {len(doc2.tables)}")

    # Count KI_* and SLOT markers
    ki_count = 0
    slot_count = 0
    for p in doc2.paragraphs:
        text = p.text
        ki_count += text.count("KI_")
        slot_count += text.count("[[SLOT_")
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                ki_count += text.count("KI_")
                slot_count += text.count("[[SLOT_")

    print(f"KI_* placeholders found: {ki_count}")
    print(f"[[SLOT_*]] markers found: {slot_count}")

    # Verify key headings still exist
    headings = [(p.style.name, p.text) for p in doc2.paragraphs
                if p.style and p.style.name.startswith("Heading")]
    print(f"Headings preserved: {len(headings)}")
    for style, text in headings:
        print(f"  {style}: {text[:80]}")


if __name__ == "__main__":
    transform_template()
