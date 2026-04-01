#!/usr/bin/env python3
"""
Rebuild Gutachten templates based on 3 real finalized Gutachten.

Replaces generic [...] placeholders with:
1. Standard legal boilerplate (identical across all real Gutachten)
2. Descriptive [[SLOT_NNN: context]] markers for AI filling
3. KI_* field references for extracted data

Run: python3 scripts/rebuild-templates.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re
import os
import sys

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'gutachtenvorlagen')

# Change marker prefix — added to every modified paragraph so the user can find/review changes
# Search for "⚡KI:" in Word to find all AI-modified spots
CHANGE_MARKER = "⚡KI: "

# ─── Standard boilerplate text from real Gutachten ────────────────────────────

VID_ERKLAERUNG = """Danach ist der gerichtlich bestellte Insolvenzverwalter oder Sachwalter in seiner Funktion im Rahmen des Insolvenzverfahrens der unabhängige, objektive, zur Sachlichkeit verpflichtete geschäftskundige Wahrer der Interessen aller im Insolvenzverfahren Beteiligten. Die erforderliche Unabhängigkeit des Unterzeichners und der geforderte Qualitätsstandard sind durch die Zertifizierung nach ISO 9001:2015 und GOI (Grundsätze ordnungsgemäßer Insolvenzverwaltung nach VID) gegeben. Die Zertifizierung wurde zuletzt am 15.07.2025 erneuert."""

VID_UNABHAENGIGKEIT = """Eine Vorbefassung, die die Unabhängigkeit des Unterzeichners in Frage stellen könnte, liegt nicht vor. Insbesondere bestand vor Antragstellung weder ein Mandatsverhältnis zur Antragstellerin noch eine sonstige beratende oder vertretende Tätigkeit."""

ZAHLUNGSUNFAEHIGKEIT_DEF = """Nach der Leitentscheidung des BGH vom 24.05.2005 (Az. IX ZR 123/04) liegt Zahlungsunfähigkeit im Sinne des § 17 InsO regelmäßig jedenfalls dann vor, wenn der Schuldner 10 % oder mehr seiner fälligen Gesamtverbindlichkeiten länger als drei Wochen nicht erfüllen kann, sofern nicht ausnahmsweise mit an Sicherheit grenzender Wahrscheinlichkeit zu erwarten ist, dass die Liquiditätslücke demnächst vollständig oder fast vollständig beseitigt werden wird und den Gläubigern ein Zuwarten nach den besonderen Umständen des Einzelfalls zuzumuten ist. Nach einer weiteren Grundsatzentscheidung des BGH (Beschl. v. 19.07.2007 – IX ZB 36/07) sind nur die „ernsthaft eingeforderten" fälligen Forderungen bei Feststellung der Zahlungsunfähigkeit zu berücksichtigen."""

UEBERSCHULDUNG_DEF_JP = """Bei juristischen Personen und diesen gleichgestellten kapitalistisch organisierten Personengesellschaften ist auch die Überschuldung ein Eröffnungsgrund. Überschuldung liegt vor, wenn das Vermögen des Schuldners die bestehenden Verbindlichkeiten nicht mehr deckt, es sei denn, die Fortführung des Unternehmens in den nächsten zwölf Monaten ist nach den Umständen überwiegend wahrscheinlich (§ 19 Abs. 2 S. 1 InsO)."""

UEBERSCHULDUNG_NP = """Überschuldung ist nur bei juristischen Personen und diesen gleichgestellten kapitalistisch organisierten Personengesellschaften ein Eröffnungsgrund. Im hiesigen Gutachtenfall ist der Insolvenzgrund unbeachtlich."""

ZUSTAENDIGKEIT_BOILERPLATE = """Die Bestimmung der örtlichen Zuständigkeit richtet sich gem. § 3 Abs. 1 S. 2 InsO in erster Linie nach dem Mittelpunkt der selbstständigen wirtschaftlichen Tätigkeit des Schuldners. Nur wenn keine selbstständige wirtschaftliche Tätigkeit mehr ausgeübt wird, ist gemäß § 3 Abs. 1 S. 1 InsO auf den allgemeinen Gerichtsstand des Schuldners abzustellen. Der Mittelpunkt der selbstständigen Tätigkeit liegt dort, wo unmittelbar Geschäfte abgeschlossen werden. Abzustellen ist auf den Ort der tatsächlichen Willensbildung, das heißt von wo aus die unternehmerischen Leitentscheidungen getroffen und in laufende Geschäftsführungsakte umgesetzt werden (OLG Brandenburg, ZInsO 2002, 767; AG Essen, ZInsO 2009, 2207)."""

BEWERTUNGSMASSSTAB_TEXT = """KI_Verwalter_Der_Die_Groß KI_Verwalter_Unterzeichner setzt bei der Bewertung des beweglichen Sachanlagevermögens Liquidations- und Fortführungswerte an."""

GRUNDLAGEN_ALLGEMEIN = """Dieses Gutachten basiert auf der Auswertung der Gerichtsakten, Gesprächen mit den unten genannten Auskunftspersonen, Inaugenscheinnahmen, den Ergebnissen des beauftragten Sachverständigen und den von KI_Schuldner_dem_der KI_Schuldner_Schuldnerin vorgelegten Unterlagen."""


def replace_paragraph_text(para, old_text, new_text):
    """Replace text in a paragraph while trying to preserve formatting."""
    if old_text not in para.text:
        return False
    # Simple case: single run
    if len(para.runs) == 1:
        para.runs[0].text = para.runs[0].text.replace(old_text, new_text)
        return True
    # Complex case: text spans multiple runs - rebuild
    full_text = para.text
    if old_text in full_text:
        new_full = full_text.replace(old_text, new_text)
        # Clear all runs except first, put all text in first run
        if para.runs:
            fmt = para.runs[0].font  # preserve first run's formatting
            for i in range(len(para.runs) - 1, 0, -1):
                para.runs[i].text = ''
            para.runs[0].text = new_full
            return True
    return False


def set_paragraph_text(para, new_text, mark=True):
    """Set entire paragraph text, preserving formatting of first run.
    If mark=True, prepend CHANGE_MARKER so user can find all changes in Word."""
    marked = (CHANGE_MARKER + new_text) if (mark and new_text.strip()) else new_text
    if para.runs:
        para.runs[0].text = marked
        for i in range(1, len(para.runs)):
            para.runs[i].text = ''
    else:
        para.add_run(marked)


def add_paragraph_after(doc, ref_para, text, style=None):
    """Add a new paragraph after a reference paragraph."""
    new_para = deepcopy(ref_para._element)
    # Clear the content
    for child in list(new_para):
        if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
            new_para.remove(child)
    ref_para._element.addnext(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, ref_para._parent)
    p.add_run(text)
    if style:
        p.style = doc.styles[style]
    return p


def process_juristische_person(template_path, output_path):
    """Update juristische Person template based on real Gutachten (freiraum 3 GmbH, Carl Puricelli)."""
    doc = Document(template_path)
    slot_counter = [1]

    def next_slot(desc):
        n = slot_counter[0]
        slot_counter[0] += 1
        return f'[[SLOT_{n:03d}: {desc}]]'

    changes = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # ─── Cover page fixes ─────────────────────────────────────────
        if text == '[juristische Person]':
            set_paragraph_text(para, '')
            changes += 1

        # ─── Ergebnis summary on page 2 ──────────────────────────────
        elif text == '[Ergebnis 1: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 1: Kernaussage zu Eröffnungsgründen (§ 17 / § 19 InsO)'))
            changes += 1
        elif text == '[Ergebnis 2: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 2: Kernaussage zur Verfahrenskostendeckung (§ 54 InsO)'))
            changes += 1
        elif text == '[Ergebnis 3: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 3: Kernaussage zu Sicherungsmaßnahmen'))
            changes += 1
        elif text == '[ggf. …]':
            set_paragraph_text(para, next_slot('Ggf. weitere Ergebnisse (Sanierung, Eigenverwaltung, Insolvenzreife)'))
            changes += 1

        # ─── A. Auftrag ───────────────────────────────────────────────
        # The Auftrag paragraph with [...] for dates and descriptions
        elif 'hat KI_Verwalter_den_die KI_Verwalter_Unterzeichner mit Beschluss vom KI_Akte_BeschlussDat' in text:
            # This is the main Auftrag paragraph - keep as is, it already uses KI_ fields
            pass

        # ─── A.II Grundlagen ──────────────────────────────────────────
        elif text.startswith('Dieses Gutachten basiert auf der Auswertung der Gerichtsakten'):
            set_paragraph_text(para, 'Dieses Gutachten basiert auf der Auswertung der Gerichtsakten, Gesprächen mit den unten genannten Auskunftspersonen, Inaugenscheinnahmen, den Ergebnissen der beauftragten Sachverständigen und den von KI_Schuldner_der_die KI_Schuldner_Schuldnerin vorgelegten Unterlagen.')
            changes += 1

        # Unterlagen section
        elif 'hat das Insolvenzgericht KI_Verwalter_dem_der KI_Verwalter_Unterzeichner die Gerichtsakte übersandt' in text:
            set_paragraph_text(para, next_slot('Datum Übersendung Gerichtsakte') + ' hat das Insolvenzgericht KI_Verwalter_dem_der KI_Verwalter_Unterzeichner die Gerichtsakte übersandt.')
            changes += 1

        elif 'fand ein erster Besprechungstermin zwischen' in text:
            set_paragraph_text(para, next_slot('Beschreibung des ersten Besprechungstermins: Datum, Teilnehmer, Ort, besprochene Themen'))
            changes += 1

        elif text == '[Angaben zu weiteren Auskunftspersonen (StB, etc.)]':
            set_paragraph_text(para, next_slot('Liste der Auskunftspersonen mit Name, Funktion und Organisation'))
            changes += 1

        # ─── A.III VID Erklärung ──────────────────────────────────────
        elif 'fand [auf Wunsch des' in text or ('Informationsgespräch' in text and '[' in text):
            set_paragraph_text(para, VID_ERKLAERUNG)
            changes += 1

        # ─── B.II Angaben zum Verfahren ───────────────────────────────
        elif text.startswith('Am […] hat […] einen Antrag'):
            set_paragraph_text(para, next_slot('Verfahrenshistorie: Wer hat wann welchen Antrag gestellt, Beschlüsse des Gerichts, Bestellung des Verwalters/Sachwalters'))
            changes += 1

        # EuInsVO section
        elif 'Art. 3 Abs. 1 EuInsVO sind deutsche Gerichte' in text:
            # Keep but note this is conditional - remove if no international connection
            pass

        # Zuständigkeit boilerplate
        elif 'Die Bestimmung der örtlichen Zuständigkeit richtet sich' in text:
            set_paragraph_text(para, ZUSTAENDIGKEIT_BOILERPLATE)
            changes += 1

        # Sicherungsmaßnahmen
        elif 'ordnete das Insolvenzgericht zur Sicherung' in text:
            set_paragraph_text(para, next_slot('Sicherungsmaßnahmen: Beschluss des Insolvenzgerichts mit Datum und Inhalt (Zustimmungsvorbehalt, Zwangsvollstreckungsverbot, Kontensperre etc.)'))
            changes += 1

        # Einzelermächtigungen
        elif 'wurde KI_Verwalter_der_die vorläufige/r Insolvenzverwalter/in ermächtigt' in text:
            set_paragraph_text(para, next_slot('Einzelermächtigungen: Welche Ermächtigungen wurden erteilt (Insolvenzgeldvorfinanzierung, Masseverbindlichkeiten, Vergütung etc.)'))
            changes += 1

        # ─── C. Informationen über das Unternehmen ────────────────────

        # Unternehmensgegenstand
        elif text == '[nicht nur Wiedergabe des oftmals wenig sagenden eingetragenen Gegenstands, sondern tatsächliche Tätigkeit]' or \
             (text.startswith('[nicht nur Wiedergabe') and 'eingetragene' in text):
            set_paragraph_text(para, next_slot('Unternehmensgegenstand: Tatsächliche Geschäftstätigkeit der Schuldnerin (nicht nur HR-Eintrag), Branche, Produkte/Dienstleistungen, Standorte'))
            changes += 1

        # Gesellschaftsrechtliche Verhältnisse
        elif text.startswith('[…Gründung') or (text.startswith('[') and 'Gründung' in text and 'Kapitalaufb' in text):
            set_paragraph_text(para, next_slot('Gesellschaftsrechtliche Verhältnisse: Gründung, Umfirmierungen, Gesellschafter (Tabelle mit Lfd.Nr, Name, Sitz, Beteiligung), Geschäftsführer, Prokurist, Stammkapital'))
            changes += 1

        # Finanzierungsstruktur
        elif text.startswith('[verbesserte Gliederungsmöglichkeit') or \
             (text.startswith('[') and 'Fremdfinanzierung' in text):
            set_paragraph_text(para, next_slot('Einleitung Finanzierungsstruktur: Überblick über die Finanzierungssituation'))
            changes += 1
        elif text.startswith('[wesentliche Finanzierungsverhältnisse'):
            set_paragraph_text(para, next_slot('Finanzierungsverhältnisse: Für jeden Kreditgeber: Name, Adresse, Art des Kredits, Betrag, Konditionen, Sicherheiten, aktueller Stand'))
            changes += 1

        # Arbeitsrechtliche Verhältnisse
        elif '[wenn überschaubare arbeitsrechtliche Verhältnisse' in text:
            # This is a heading with editorial instruction - clean it
            pass  # heading text handled separately
        elif 'beschäftigt […] Arbeitnehmer, davon' in text and 'Ausbildung' in text:
            set_paragraph_text(para, 'KI_Schuldner_Der_Die_Groß KI_Schuldner_Schuldnerin beschäftigt KI_Arbeitnehmer_Anzahl Arbeitnehmer. ' + next_slot('Arbeitsrechtliche Details: Anzahl Auszubildende, Betriebsrat (ja/nein), Lohnrückstände, Insolvenzgeldvorfinanzierung'))
            changes += 1
        elif text == '[wenn größerer Betrieb]':
            set_paragraph_text(para, '')
            changes += 1

        # Vertragsverhältnisse descriptions
        elif text == '[Gewerberaummiet-/Pachtverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Gewerberaummietverträge: Vermieter, Adresse, Mietfläche, monatliche Miete, Laufzeit, Rückstände'))
            changes += 1
        elif text == '[Versicherungsverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Versicherungsverträge: Art, Versicherungsnehmer, Versicherungsnummer, Sicherungsabtretungen'))
            changes += 1
        elif text == '[Versorgungsverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Versorgungsverträge: Strom, Gas, Wasser, Telekommunikation'))
            changes += 1
        elif text.startswith('[… Finanzierungs-Leasing]') or text == '[… Finanzierungs-Leasing]':
            set_paragraph_text(para, next_slot('Leasingverträge: Leasinggeber, Gegenstand, monatliche Rate, Vertragsnummer, Laufzeit'))
            changes += 1
        elif text == '[Factoring-Vereinbarungen]':
            set_paragraph_text(para, next_slot('Factoring: Factor, Konditionen, aktuelles Volumen'))
            changes += 1
        elif text == '[sonstige Vertragsverhältnisse]':
            set_paragraph_text(para, next_slot('Sonstige Verträge: Wartung, IT, Beratung etc.'))
            changes += 1

        # Zahlungsverkehr
        elif 'ggf. Sicherheiten nennen' in text:
            set_paragraph_text(para, next_slot('Zahlungsverkehr: Bankverbindungen mit Kontonummern und Salden bei Antragstellung, Insolvenzsonderkonto'))
            changes += 1

        # Rechnungswesen
        elif text.startswith('Die Buchführung') and '[letzter Jahresabschluss' in text:
            set_paragraph_text(para, next_slot('Buchhaltung vor Antragstellung: Steuerberater, letzte Steuererklärungen, letzter Jahresabschluss, Buchführungsrückstände'))
            changes += 1
        elif '[…]' == text and i > 0 and 'Buchhaltung nach' in doc.paragraphs[i-2].text if i >= 2 else False:
            set_paragraph_text(para, next_slot('Buchhaltung nach Antragstellung: Wer führt die insolvenzspezifische Buchhaltung'))
            changes += 1

        # Anhängige Rechtsstreitigkeiten
        elif text == '[anhängige Rechtsstreitigkeiten]':
            set_paragraph_text(para, next_slot('Anhängige Rechtsstreitigkeiten: Bekannte Klagen und deren Status, oder "Anhängige Rechtsstreitigkeiten sind dem Unterzeichner nicht bekannt."'))
            changes += 1

        # Betriebswirtschaftliche Verhältnisse
        elif text.startswith('[ggf. North-Data-Daten'):
            set_paragraph_text(para, next_slot('Wirtschaftliche Entwicklung: Umsatzentwicklung, Krisenursachen, wirtschaftlicher Hintergrund'))
            changes += 1
        elif text.startswith('[Inventarisierung'):
            set_paragraph_text(para, next_slot('Maßnahmen im Eröffnungsverfahren: Datenraum, Datensicherung, Sicherung des Lebensunterhalts, Vermögenssicherung, Inventarisierung'))
            changes += 1
        elif text.startswith('[Zur Insolvenzgeldvorfinanzierung:') or text == '[Zur Insolvenzgeldvorfinanzierung:]':
            set_paragraph_text(para, next_slot('Insolvenzgeldvorfinanzierung: Rahmenvereinbarung, Bank, Zeitraum, Durchführung'))
            changes += 1
        elif text.startswith('[zur Kommunikation'):
            set_paragraph_text(para, next_slot('Kommunikation mit Stakeholdern: Betriebsversammlungen, Gläubigergespräche'))
            changes += 1
        elif text.startswith('[zur Sicherung der Liquidität'):
            set_paragraph_text(para, next_slot('Liquiditätssicherung: Maßnahmen zur Aufrechterhaltung der Zahlungsfähigkeit'))
            changes += 1
        elif text.startswith('[zur Aufrechterhaltung von Lieferbeziehungen'):
            set_paragraph_text(para, next_slot('Lieferbeziehungen: Maßnahmen zur Sicherung wichtiger Lieferanten'))
            changes += 1
        elif text.startswith('[Etablierung eines Systems'):
            set_paragraph_text(para, next_slot('Kontrollsysteme: Bestell- und Zahlungsfreigaben, Berichtswesen'))
            changes += 1
        elif text.startswith('[Einleitung eines Investorenprozesses'):
            set_paragraph_text(para, next_slot('Sanierungsaussichten: Investorenprozess, Fortführungsperspektive, oder "Es besteht keine Aussicht auf Sanierung"'))
            changes += 1

        # ─── D. Vermögensverhältnisse ─────────────────────────────────

        # Bewertungsmaßstab
        elif text.startswith('[Angabe ob nur Liquidations-'):
            set_paragraph_text(para, 'KI_Verwalter_Der_Die_Groß KI_Verwalter_Unterzeichner setzt bei der Bewertung des beweglichen Sachanlagevermögens Liquidations- und Fortführungswerte an.')
            changes += 1

        # Sachverständige Bewertung
        elif i > 0 and 'Sachverständige Bewertung' in doc.paragraphs[max(0,i-3):i+1][-1].text if len(doc.paragraphs) > i else False:
            pass  # handled below

        # Aktiva subsections
        elif text == '[Das Unternehmen hat einen handelsrechtlich nicht bilanzierbaren Goodwill/Firmenwert, der sich aus der Geschäftstätigkeit ergibt]' or \
             (text.startswith('[Das Unternehmen hat einen') and 'Goodwill' in text):
            set_paragraph_text(para, next_slot('Goodwill/Firmenwert: Bewertung des immateriellen Geschäftswerts, Kunden-/Lieferantenbeziehungen, Know-how'))
            changes += 1
        elif text == '[selbst geschaffene Schutzrechte Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Selbst geschaffene Schutzrechte: Patente, Marken, Beschreibung und Wert'))
            changes += 1
        elif text == '[entgeltlich erworbene Rechte Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Entgeltlich erworbene Rechte: Lizenzen, Software, Beschreibung und Wert'))
            changes += 1
        elif text == '[Grundstücke Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Grundstücke und grundstücksgleiche Rechte: Grundbuch, Gemarkung, Flurstück, Belastungen Abt. II und III, Bewertung'))
            changes += 1
        elif text == '[technische Anlagen Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Technische Anlagen und Maschinen: Auflistung mit Fabrikat, Typ, Baujahr, Betriebsstunden, Sicherungsrechte, Bewertung'))
            changes += 1
        elif text == '[BGA Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Betriebs- und Geschäftsausstattung: Büroausstattung, IT, Werkzeuge, Bewertung'))
            changes += 1
        elif text == '[Anteile verbundene Unternehmen Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Anteile an verbundenen Unternehmen: Gesellschaft, Beteiligung, Bewertung'))
            changes += 1
        elif text == '[Beteiligungen Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Beteiligungen: Gesellschaft, Anteil, Bewertung'))
            changes += 1
        elif text == '[Wertpapiere Anlagevermögen Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Wertpapiere des Anlagevermögens: Art, Nennwert, Bewertung'))
            changes += 1
        elif text == '[RHB-Stoffe Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Roh-, Hilfs- und Betriebsstoffe: Bestandsaufnahme, Beschreibung, Bewertung'))
            changes += 1
        elif text == '[unfertige Erzeugnisse Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Unfertige Erzeugnisse/Leistungen: Beschreibung, Fertigstellungsgrad, Bewertung'))
            changes += 1
        elif text == '[fertige Erzeugnisse/Waren Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Fertige Erzeugnisse und Waren: Bestand, Beschreibung, Bewertung'))
            changes += 1

        # Kapitalansprüche (GmbH-spezifisch)
        elif text.startswith('[bei GmbH §') or text.startswith('[Ansprüche auf Einzahlung'):
            set_paragraph_text(para, next_slot('Ansprüche auf Einzahlung der Stammeinlagen gem. § 19 Abs. 1 GmbHG'))
            changes += 1
        elif text.startswith('[Ansprüche auf Erstattung verbotener') or 'verbotener Rückzahlungen' in text:
            set_paragraph_text(para, next_slot('Ansprüche auf Erstattung verbotener Rückzahlungen gem. §§ 31 Abs. 1, 30 GmbHG'))
            changes += 1

        # ─── E. Eröffnungsgründe ──────────────────────────────────────

        # § 17 Definition - replace with standard boilerplate
        elif 'Leitentscheidung des BGH' in text and '24.05.2005' in text:
            # Already has the right text or close to it
            set_paragraph_text(para, ZAHLUNGSUNFAEHIGKEIT_DEF)
            changes += 1

        # § 19 Definition for juristische Person
        elif 'juristischen Personen' in text and 'Überschuldung ein Eröffnungsgrund' in text:
            set_paragraph_text(para, UEBERSCHULDUNG_DEF_JP)
            changes += 1

        # ─── F. Insolvenzspezifische Ansprüche ────────────────────────

        # ─── G. Verfahrenskostendeckung ───────────────────────────────

        # ─── H. Ergebnis ──────────────────────────────────────────────

        # ─── Generic [...] cleanup ────────────────────────────────────
        # Replace remaining standalone [...] and [...]
        elif text == '[…]' or text == '[...]':
            # Look at context from nearby headings
            context = ''
            for j in range(max(0, i-5), i):
                prev = doc.paragraphs[j].text.strip()
                if prev and doc.paragraphs[j].style.name.startswith('Heading'):
                    context = prev
                    break
            if context:
                set_paragraph_text(para, next_slot(f'{context}: Fallspezifische Angaben'))
            else:
                set_paragraph_text(para, next_slot('Fallspezifische Angaben'))
            changes += 1

    # ─── Clean up heading editorial instructions ──────────────────────
    for para in doc.paragraphs:
        text = para.text
        if para.style.name.startswith('Heading'):
            # Remove editorial instructions in brackets from headings
            cleaned = re.sub(r'\s*\[.*?\]\s*$', '', text).strip()
            if cleaned != text:
                set_paragraph_text(para, cleaned)
                changes += 1
            # Fix [...] in heading text
            if '[…]' in text or '[...]' in text:
                set_paragraph_text(para, text.replace('[…]', 'KI_Gericht_Ort').replace('[...]', 'KI_Gericht_Ort'))
                changes += 1

    # ─── Handle remaining [...] in body text ──────────────────────────
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not para.style.name.startswith('Heading') and ('[…]' in text or '[...]' in text):
            # Replace inline [...] with contextual slot markers
            # Count how many [...] in this paragraph
            count = text.count('[…]') + text.count('[...]')
            if count <= 3:
                new_text = text
                for _ in range(count):
                    slot = next_slot('Angabe')
                    new_text = new_text.replace('[…]', slot, 1).replace('[...]', slot, 1)
                set_paragraph_text(para, new_text)
                changes += 1

    # ─── Handle remaining [description] brackets ──────────────────────
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not para.style.name.startswith('Heading'):
            brackets = re.findall(r'\[([^\]]{3,})\]', text)
            for bracket_content in brackets:
                if bracket_content.startswith('SLOT_') or bracket_content.startswith('KI_'):
                    continue  # Already a slot marker
                old = f'[{bracket_content}]'
                new = next_slot(bracket_content[:60])
                if old in text:
                    text = text.replace(old, new, 1)
                    changes += 1
            if text != para.text:
                set_paragraph_text(para, text)

    # ─── Also process tables ──────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if '[…]' in text or '[...]' in text:
                        new_text = text.replace('[…]', next_slot('Tabellenwert')).replace('[...]', next_slot('Tabellenwert'))
                        set_paragraph_text(para, new_text)
                        changes += 1
                    brackets = re.findall(r'\[([^\]]{3,})\]', text)
                    for bc in brackets:
                        if bc.startswith('SLOT_') or bc.startswith('KI_'):
                            continue
                        old = f'[{bc}]'
                        new = next_slot(bc[:40])
                        if old in para.text:
                            set_paragraph_text(para, para.text.replace(old, new, 1))
                            changes += 1

    print(f'  Applied {changes} changes, {slot_counter[0]-1} slots created')
    doc.save(output_path)
    print(f'  Saved: {output_path}')


def process_natuerliche_person(template_path, output_path):
    """Update natürliche Person template based on real Gutachten (Alexander Geldt)."""
    doc = Document(template_path)
    slot_counter = [1]

    def next_slot(desc):
        n = slot_counter[0]
        slot_counter[0] += 1
        return f'[[SLOT_{n:03d}: {desc}]]'

    changes = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # ─── Ergebnis summary on page 2 ──────────────────────────────
        if text == '[Ergebnis 1: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 1: Kernaussage zu Eröffnungsgründen (§ 17 InsO)'))
            changes += 1
        elif text == '[Ergebnis 2: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 2: Kernaussage zur Verfahrenskostendeckung (§ 54 InsO)'))
            changes += 1
        elif text == '[Ergebnis 3: Kernaussage Gutachten]':
            set_paragraph_text(para, next_slot('Ergebnis 3: Kernaussage zu Sicherungsmaßnahmen'))
            changes += 1
        elif text == '[ggf. …]':
            set_paragraph_text(para, next_slot('Ggf. weitere Ergebnisse (Insolvenzgeldzeitraum, Empfehlung Eröffnungsdatum)'))
            changes += 1

        # ─── Grundlagen ───────────────────────────────────────────────
        elif text.startswith('Dieses Gutachten basiert auf der Auswertung der Gerichtsakten'):
            set_paragraph_text(para, 'Dieses Gutachten basiert auf der Auswertung der Gerichtsakten, Gesprächen mit den unten genannten Auskunftspersonen, Inaugenscheinnahmen, den Ergebnissen des beauftragten Sachverständigen und den von KI_Schuldner_dem_der KI_Schuldner_Schuldnerin vorgelegten Unterlagen.')
            changes += 1

        elif 'hat das Insolvenzgericht KI_Verwalter_dem_der KI_Verwalter_Unterzeichner die Gerichtsakte übersandt' in text:
            set_paragraph_text(para, next_slot('Datum Übersendung Gerichtsakte') + ' hat das Insolvenzgericht KI_Verwalter_dem_der KI_Verwalter_Unterzeichner die Gerichtsakte übersandt.')
            changes += 1

        elif 'fand ein erster Besprechungstermin zwischen' in text:
            set_paragraph_text(para, next_slot('Beschreibung des ersten Besprechungstermins: Datum, Teilnehmer, Ort, Unterlagen angefordert'))
            changes += 1

        elif text == '[Angaben zu weiteren Auskunftspersonen (StB, etc.)]':
            set_paragraph_text(para, next_slot('Liste der Auskunftspersonen mit Name und Organisation'))
            changes += 1

        # ─── VID Erklärung ────────────────────────────────────────────
        elif 'Informationsgespräch' in text and '[' in text:
            set_paragraph_text(para, VID_ERKLAERUNG)
            changes += 1

        # ─── Verfahren ────────────────────────────────────────────────
        elif text.startswith('Am […] hat') and 'Antrag' in text:
            set_paragraph_text(para, next_slot('Verfahrenshistorie: Antragstellung, Beschlüsse, Bestellung des Sachverständigen/Verwalters'))
            changes += 1

        elif 'Die Bestimmung der örtlichen Zuständigkeit richtet sich' in text:
            set_paragraph_text(para, ZUSTAENDIGKEIT_BOILERPLATE)
            changes += 1

        elif 'ordnete das Insolvenzgericht zur Sicherung' in text:
            set_paragraph_text(para, next_slot('Sicherungsmaßnahmen des Insolvenzgerichts mit Datum und Inhalt'))
            changes += 1

        elif 'wurde KI_Verwalter_der_die vorläufige' in text and 'ermächtigt' in text:
            set_paragraph_text(para, next_slot('Einzelermächtigungen: Insolvenzgeldvorfinanzierung, Masseverbindlichkeiten etc.'))
            changes += 1

        # ─── Informationen ────────────────────────────────────────────
        elif text.startswith('[nicht nur Wiedergabe') or (text.startswith('[') and 'eingetragene' in text and 'Gegenstand' in text):
            set_paragraph_text(para, next_slot('Persönliche Verhältnisse und Unternehmensgegenstand: Geburtsdatum, Familienstand, Firmenbeschreibung, tatsächliche Tätigkeit'))
            changes += 1

        elif text.startswith('[wesentliche Finanzierungsverhältnisse'):
            set_paragraph_text(para, next_slot('Finanzierungsstruktur: Für jede Bank: Konten, Salden, Sicherheiten'))
            changes += 1

        elif '[wenn überschaubare arbeitsrechtliche Verhältnisse' in text:
            pass  # heading - handled below

        elif 'beschäftigt […] Arbeitnehmer, davon' in text:
            set_paragraph_text(para, 'KI_Schuldner_Der_Die_Groß KI_Schuldner_Schuldnerin beschäftigt KI_Arbeitnehmer_Anzahl Arbeitnehmer. ' + next_slot('Arbeitsrechtliche Details'))
            changes += 1
        elif text == '[wenn größerer Betrieb]':
            set_paragraph_text(para, '')
            changes += 1

        elif text == '[Gewerberaummiet-/Pachtverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Gewerberaummietverträge: Vermieter, Fläche, Miete, Rückstände'))
            changes += 1
        elif text == '[Versicherungsverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Versicherungsverträge: Art, Versicherungsnummer, Sicherungsabtretungen'))
            changes += 1
        elif text == '[Versorgungsverträge Beschreibung]':
            set_paragraph_text(para, next_slot('Versorgungsverträge: Strom, Gas, Wasser'))
            changes += 1
        elif text.startswith('[… Finanzierungs-Leasing]'):
            set_paragraph_text(para, next_slot('Leasingverträge und Mietkaufverträge: Gegenstand, Rate, Laufzeit'))
            changes += 1
        elif text == '[Factoring-Vereinbarungen]':
            set_paragraph_text(para, next_slot('Factoring-Vereinbarungen'))
            changes += 1
        elif text == '[sonstige Vertragsverhältnisse]':
            set_paragraph_text(para, next_slot('Sonstige Verträge'))
            changes += 1
        elif 'ggf. Sicherheiten nennen' in text:
            set_paragraph_text(para, next_slot('Zahlungsverkehr: Bankverbindungen, Konten, Salden, Insolvenzsonderkonto'))
            changes += 1
        elif '[letzter Jahresabschluss' in text:
            set_paragraph_text(para, next_slot('Buchhaltung vor Antragstellung: Steuerberater, letzte Erklärungen, Abschlüsse'))
            changes += 1
        elif text == '[anhängige Rechtsstreitigkeiten]':
            set_paragraph_text(para, next_slot('Anhängige Rechtsstreitigkeiten oder "nicht bekannt"'))
            changes += 1
        elif text.startswith('[ggf. North-Data'):
            set_paragraph_text(para, next_slot('Wirtschaftliche Entwicklung und Krisenursache'))
            changes += 1
        elif text.startswith('[Inventarisierung'):
            set_paragraph_text(para, next_slot('Maßnahmen im Eröffnungsverfahren'))
            changes += 1
        elif text.startswith('[Zur Insolvenzgeldvorfinanzierung'):
            set_paragraph_text(para, next_slot('Insolvenzgeldvorfinanzierung'))
            changes += 1
        elif text.startswith('[zur Kommunikation'):
            set_paragraph_text(para, next_slot('Kommunikation mit Stakeholdern'))
            changes += 1
        elif text.startswith('[zur Sicherung der Liquidität'):
            set_paragraph_text(para, next_slot('Liquiditätssicherung'))
            changes += 1
        elif text.startswith('[zur Aufrechterhaltung'):
            set_paragraph_text(para, next_slot('Lieferbeziehungen'))
            changes += 1
        elif text.startswith('[Etablierung'):
            set_paragraph_text(para, next_slot('Kontrollsysteme'))
            changes += 1
        elif text.startswith('[Einleitung eines Investor'):
            set_paragraph_text(para, next_slot('Sanierungsaussichten'))
            changes += 1

        # ─── Vermögensverhältnisse ────────────────────────────────────
        elif text.startswith('[Angabe ob nur Liquidations-'):
            set_paragraph_text(para, 'KI_Verwalter_Der_Die_Groß KI_Verwalter_Unterzeichner setzt bei der Bewertung des beweglichen Sachanlagevermögens Liquidations- und Fortführungswerte an.')
            changes += 1

        elif text == '[Grundstücke Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Grundstücke: Grundbuch, Gemarkung, Belastungen, Bewertung'))
            changes += 1
        elif text == '[technische Anlagen Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Technische Anlagen und Maschinen: Einzelauflistung mit Bewertung'))
            changes += 1
        elif text == '[BGA Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Betriebs- und Geschäftsausstattung: Beschreibung und Bewertung'))
            changes += 1
        elif text == '[RHB-Stoffe Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Roh-, Hilfs- und Betriebsstoffe'))
            changes += 1
        elif text == '[unfertige Erzeugnisse Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Unfertige Erzeugnisse'))
            changes += 1
        elif text == '[fertige Erzeugnisse/Waren Beschreibung/Wert]':
            set_paragraph_text(para, next_slot('Fertige Erzeugnisse und Waren'))
            changes += 1

        # ─── Eröffnungsgründe ─────────────────────────────────────────
        elif 'Leitentscheidung des BGH' in text and '24.05.2005' in text:
            set_paragraph_text(para, ZAHLUNGSUNFAEHIGKEIT_DEF)
            changes += 1

        # § 19 for natürliche Person - not applicable
        elif 'nur bei juristischen Personen' in text and 'Überschuldung' in text:
            set_paragraph_text(para, UEBERSCHULDUNG_NP)
            changes += 1

        # ─── Generic [...] cleanup ────────────────────────────────────
        elif text == '[…]' or text == '[...]':
            context = ''
            for j in range(max(0, i-5), i):
                prev = doc.paragraphs[j].text.strip()
                if prev and doc.paragraphs[j].style.name.startswith('Heading'):
                    context = prev
                    break
            if context:
                set_paragraph_text(para, next_slot(f'{context}: Angaben'))
            else:
                set_paragraph_text(para, next_slot('Angaben'))
            changes += 1

    # ─── Clean headings ───────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text
        if para.style.name.startswith('Heading'):
            cleaned = re.sub(r'\s*\[.*?\]\s*$', '', text).strip()
            if cleaned != text:
                set_paragraph_text(para, cleaned)
                changes += 1
            if '[…]' in text or '[...]' in text:
                set_paragraph_text(para, text.replace('[…]', 'KI_Gericht_Ort').replace('[...]', 'KI_Gericht_Ort'))
                changes += 1

    # ─── Remaining inline [...] ───────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text
        if not para.style.name.startswith('Heading') and ('[…]' in text or '[...]' in text):
            count = text.count('[…]') + text.count('[...]')
            if count <= 5:
                new_text = text
                for _ in range(count):
                    slot = next_slot('Angabe')
                    new_text = new_text.replace('[…]', slot, 1).replace('[...]', slot, 1)
                set_paragraph_text(para, new_text)
                changes += 1

    # ─── Remaining [description] brackets ─────────────────────────────
    for para in doc.paragraphs:
        text = para.text
        if not para.style.name.startswith('Heading'):
            brackets = re.findall(r'\[([^\]]{3,})\]', text)
            for bc in brackets:
                if bc.startswith('SLOT_') or bc.startswith('KI_'):
                    continue
                old = f'[{bc}]'
                new = next_slot(bc[:60])
                if old in text:
                    text = text.replace(old, new, 1)
                    changes += 1
            if text != para.text:
                set_paragraph_text(para, text)

    # ─── Tables ───────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if '[…]' in text or '[...]' in text:
                        new_text = text.replace('[…]', next_slot('Tabellenwert')).replace('[...]', next_slot('Tabellenwert'))
                        set_paragraph_text(para, new_text)
                        changes += 1
                    brackets = re.findall(r'\[([^\]]{3,})\]', text)
                    for bc in brackets:
                        if bc.startswith('SLOT_') or bc.startswith('KI_'):
                            continue
                        old = f'[{bc}]'
                        new = next_slot(bc[:40])
                        if old in para.text:
                            set_paragraph_text(para, para.text.replace(old, new, 1))
                            changes += 1

    print(f'  Applied {changes} changes, {slot_counter[0]-1} slots created')
    doc.save(output_path)
    print(f'  Saved: {output_path}')


def verify_template(path):
    """Verify template can be opened and count remaining issues."""
    doc = Document(path)
    remaining_brackets = 0
    ki_fields = set()
    slots = set()
    for para in doc.paragraphs:
        text = para.text
        remaining_brackets += len(re.findall(r'\[[^\]]*…[^\]]*\]', text))
        remaining_brackets += len(re.findall(r'\[\.\.\.\]', text))
        for m in re.findall(r'KI_\w+', text):
            ki_fields.add(m)
        for m in re.findall(r'\[\[SLOT_\d+:', text):
            slots.add(m)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    remaining_brackets += len(re.findall(r'\[[^\]]*…[^\]]*\]', text))
                    for m in re.findall(r'KI_\w+', text):
                        ki_fields.add(m)
                    for m in re.findall(r'\[\[SLOT_\d+:', text):
                        slots.add(m)
    print(f'  Remaining [...]: {remaining_brackets}')
    print(f'  KI_ fields: {len(ki_fields)}')
    print(f'  [[SLOT_*]] markers: {len(slots)}')
    print(f'  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
    return remaining_brackets


if __name__ == '__main__':
    print('=== Rebuilding Gutachten templates ===\n')

    # Juristische Person
    jp_in = os.path.join(TEMPLATES_DIR, 'Gutachten Muster juristische Person.docx')
    jp_out = os.path.join(TEMPLATES_DIR, 'Gutachten Muster juristische Person.docx')
    print(f'Processing: juristische Person')
    process_juristische_person(jp_in, jp_out)
    print(f'Verifying:')
    verify_template(jp_out)
    print()

    # Natürliche Person
    np_in = os.path.join(TEMPLATES_DIR, 'Gutachten Muster natürliche Person.docx')
    np_out = os.path.join(TEMPLATES_DIR, 'Gutachten Muster natürliche Person.docx')
    print(f'Processing: natürliche Person')
    process_natuerliche_person(np_in, np_out)
    print(f'Verifying:')
    verify_template(np_out)
    print()

    print('Done! Templates updated in place.')
    print('Back up originals if needed.')
